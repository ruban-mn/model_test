import pandas as pd
import numpy as np
import openpyxl
import docx
import streamlit as st
import io
import time as tm
import re
from docx import Document
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import math



st.subheader('Независимая оценка качества услуг медицинских организаций')

uploaded_file = st.file_uploader("**Загрузите файл сводную по чек-листам амбулатория**", type=["xls", "xlsx"])
chek_list = pd.read_excel(uploaded_file)

uploaded_file1 = st.file_uploader("**Загрузите файл сводную по чек-листам стационар**", type=["xls", "xlsx"])
chek_list_st = pd.read_excel(uploaded_file1)

uploaded_file2 = st.file_uploader("**Загрузите файл с массивом анкет амбулатория**", type=["xls", "xlsx"])
Answers_respond = pd.read_excel(uploaded_file2)

uploaded_file3 = st.file_uploader("**Загрузите файл с массивом анкет стационар**", type=["xls", "xlsx"])
Answers_respond_st = pd.read_excel(uploaded_file3)

plase = st.text_input('Введите название территории в родительном падеже', 'н-р Московской области')

## подгружаем сводную по чек листам
##chek_list=pd.read_excel(r"C:\Users\user\чек лист гулькевичи.xlsx")
##подгружаем массив с ответами респондентов
##Answers_respond=pd.read_excel(r"C:\Users\user\Анкета Гулькевичи НОК культура (Ответы).xlsx")

tm.sleep(30)

Answers_respond_list = Answers_respond.columns.tolist() ##извлекаем наименования столбцов в список

New_col = []  # Создаем пустой список
for i in range(32):  # Цикл от 0 до 18
    sim = i   # присваиваем номер
    New_col.append('v' + str(sim))  # добавляем новый номер вопрса в список

dictionary = dict(zip(Answers_respond_list, New_col)) # создаем  словарь для переименования стобцов
Answers_respond = Answers_respond.rename(columns=dictionary) # переименовываем столбцы в начальном датафрейме
Answers_respond['v32'] = Answers_respond['v2'].map({'менее 1 часа': 1, '3 часа': 3})
Answers_respond['v33'] = Answers_respond['v7'].map({'менее 7 календарных дней': 6, '7 календарных дней': 7, '10 календарных дней': 10, '12 календарных дней': 12, '13 календарных дней': 13, '14 календарных дней и более': 14})

# Рассчитываем значение для нового столбца
def calculate_value(row):
    total_answers = sum([val for val in row['v32'] if not math.isnan(val)])

    result_sum = sum([(i+1)*val for i, val in enumerate(row['v32']) if not math.isnan(val)])
    result = round(result_sum / total_answers)
    
    if result == 24:
        return 10
    elif result == 23:
        return 20
    elif result == 22:
        return 40
    elif 21 <= result <= 13:
        return 60
    else:
        return 100


result_df = Answers_respond.groupby('v0').apply(calculate_value).reset_index()
result_df.columns = ['v0', 'v32']

# Рассчитываем значение для нового столбца
def calculate_value1(row):
    total_answers = sum([val for val in row['v33'] if not math.isnan(val)])

    result_sum = sum([(i+1)*val for i, val in enumerate(row['v33']) if not math.isnan(val)])
    result = round(result_sum / total_answers)
    
    if result == 14:
        return 10
    elif result == 13:
        return 20
    elif result == 12:
        return 40
    elif 11 <= result <= 8:
        return 60
    else:
        return 100


result_df1 = Answers_respond.groupby('v0').apply(calculate_value1).reset_index()
result_df1.columns = ['v0', 'v33']

# Создание нового DataFrame для хранения результатов подсчета, считам количество ответов да на вопросы анкеты
ans_res = pd.DataFrame({'v0': Answers_respond['v0'].unique()})

selected_columns = ['v5', 'v10', 'v11', 'v12', 'v13', 'v14', 'v15', 'v16', 'v17', 'v19', 'v23', 'v26', 'v27', 'v28', 'v29', 'v30', 'v31']
# Используем цикл для подсчета значений и создания новых столбцов
for col in selected_columns:
    value = 'да'  # Значение, которое мы считаем
    count_col_name = f'_{col}_'
    counts = Answers_respond[Answers_respond[col] == value].groupby('v0').size().reset_index(name=count_col_name)
    ans_res = ans_res.merge(counts, on='v0', how='left')

ans_res = ans_res.merge(result_df, on='v0', how='left')

ans_res = ans_res.merge(result_df1, on='v0', how='left')

ans_res = ans_res.dropna(axis=1) # Удаляем столбцы со значением NaN
ans_res['v0'] = ans_res['v0'].str.replace('.', '')# Удаляем точку из наименований организаций
ans_res = ans_res.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
ans_res = ans_res.reset_index(drop=True)

col_ob = Answers_respond.groupby('v0').size().reset_index(name='Ч_общ')
col_ob['v0'] = col_ob['v0'].str.replace('.', '')# Удаляем точку из наименований организаций
col_ob = col_ob.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
col_ob = col_ob.reset_index(drop=True)
all_ans = col_ob['Ч_общ']

name_org = chek_list.filter(like='Наименование организации').copy()

Raschet_ballov = name_org
Raschet_ballov['Истенд'] = chek_list.filter(like='На СТЕНДЕ').sum(axis=1)
Raschet_ballov['Исайт'] = chek_list.filter(like='На САЙТЕ').sum(axis=1)
Raschet_ballov['Инорм-стенд'] = chek_list.filter(like='На СТЕНДЕ').count(axis=1)
Raschet_ballov['Инорм-сайт'] = chek_list.filter(like='На САЙТЕ').count(axis=1)
Raschet_ballov['Пинф'] = round(0.5*((Raschet_ballov['Истенд']/Raschet_ballov['Инорм-стенд'])+(Raschet_ballov['Исайт']/Raschet_ballov['Инорм-сайт']))*100, 2)
Raschet_ballov['Тдист'] = 30
Raschet_ballov['Сдист'] = chek_list.filter(like='Наличие и функционирование на официальном сайте').sum(axis=1)
Raschet_ballov['Пдист'] = Raschet_ballov['Тдист']*Raschet_ballov['Сдист']
Raschet_ballov['Пдист'].where(Raschet_ballov['Пдист'] <= 100, 100, inplace=True)
Raschet_ballov['Устенд'] = ans_res['_v14_']
Raschet_ballov['Усайт'] = ans_res['_v16_']
Raschet_ballov['Уобщ-стенд'] = ans_res['_v13_']
Raschet_ballov['Уобщ-сайт'] = ans_res['_v15_']
Raschet_ballov['Поткруд'] = round(0.5*((Raschet_ballov['Устенд']/Raschet_ballov['Уобщ-стенд'])+(Raschet_ballov['Усайт']/Raschet_ballov['Уобщ-сайт']))*100, 2)
Raschet_ballov['К1'] = round(0.3*Raschet_ballov['Пинф'] + 0.3*Raschet_ballov['Пдист'] + 0.4*Raschet_ballov['Поткруд'], 2)
Raschet_ballov['Ткомф'] = chek_list.filter(like='Обеспечение в организации комфортных условий').sum(axis=1)
Raschet_ballov['Скомф'] = 20
Raschet_ballov['Пкомф.усл'] = Raschet_ballov['Ткомф']*Raschet_ballov['Скомф']
Raschet_ballov['Пкомф.усл'].where(Raschet_ballov['Пкомф.усл'] <= 100, 100, inplace=True)
Raschet_ballov['ожид'] = (ans_res['v32']+ans_res['v33'])/2
Raschet_ballov['Усвоевр'] = (ans_res['_v11_']+ans_res['_v26_'])/2
Raschet_ballov['Чобщ'] = all_ans
Raschet_ballov['Пожид'] = (round(Raschet_ballov['Усвоевр']/Raschet_ballov['Чобщ']*100, 2) + Raschet_ballov['ожид'])/2
Raschet_ballov['Укомф'] = ans_res['_v17_']
Raschet_ballov['Чобщ'] = all_ans
Raschet_ballov['Пкомфуд'] = round(Raschet_ballov['Укомф']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['К2'] = round(0.3*Raschet_ballov['Пкомф.усл'] + 0.4*Raschet_ballov['Пожид'] + 0.3*Raschet_ballov['Пкомфуд'], 2)
Raschet_ballov['Торгдост'] = chek_list.filter(like='Оборудование территории').sum(axis=1)
Raschet_ballov['Соргдост'] = 20
Raschet_ballov['Поргдост'] = Raschet_ballov['Торгдост']*Raschet_ballov['Соргдост']
Raschet_ballov['Поргдост'].where(Raschet_ballov['Поргдост'] <= 100, 100, inplace=True)
Raschet_ballov['Туслугдост'] = chek_list.filter(like='Обеспечение в организации условий доступности').sum(axis=1)
Raschet_ballov['Суслугдост'] = 20
Raschet_ballov['Пуслугдост'] = Raschet_ballov['Туслугдост']*Raschet_ballov['Суслугдост']
Raschet_ballov['Пуслугдост'].where(Raschet_ballov['Пуслугдост'] <= 100, 100, inplace=True)
Raschet_ballov['Удост'] = ans_res['_v23_']
Raschet_ballov['Чинв'] = ans_res['_v19_']
Raschet_ballov['Пдостуд'] = round(Raschet_ballov['Удост']/Raschet_ballov['Чинв']*100, 2)
Raschet_ballov['К3'] = round(0.3*Raschet_ballov['Поргдост'] + 0.4*Raschet_ballov['Пуслугдост'] + 0.3*Raschet_ballov['Пдостуд'], 2)
Raschet_ballov['Уперв.конт'] = (ans_res['_v5_'] + ans_res['_v10_'])/2
Raschet_ballov['Чобщ1'] = all_ans
Raschet_ballov['Пперв.контуд'] = round(Raschet_ballov['Уперв.конт']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['Уоказ.услуг'] = ans_res['_v12_']
Raschet_ballov['Чобщ2'] = all_ans
Raschet_ballov['Показ.услугуд'] = round(Raschet_ballov['Уоказ.услуг']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['Увежл.дист'] = ans_res['_v31_']
Raschet_ballov['Чобщ_ус'] = ans_res['_v30_']
Raschet_ballov['Пвежл.дистуд'] = round(Raschet_ballov['Увежл.дист']/Raschet_ballov['Чобщ_ус']*100, 2)
Raschet_ballov['К4'] = round(0.4*Raschet_ballov['Пперв.контуд'] + 0.4*Raschet_ballov['Показ.услугуд'] + 0.2*Raschet_ballov['Пвежл.дистуд'], 2)
Raschet_ballov['Уреком'] = ans_res['_v27_']
Raschet_ballov['Чобщ3'] = all_ans
Raschet_ballov['Преком'] = round(Raschet_ballov['Уреком']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['Уорг.усл'] = ans_res['_v28_']
Raschet_ballov['Чобщ4'] = all_ans
Raschet_ballov['Порг.услуд'] = round(Raschet_ballov['Уорг.усл']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['Ууд'] = ans_res['_v29_']
Raschet_ballov['Чобщ5'] = all_ans
Raschet_ballov['Пуд'] = round(Raschet_ballov['Ууд']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['К5'] = round(0.3*Raschet_ballov['Преком'] + 0.2*Raschet_ballov['Порг.услуд'] + 0.5*Raschet_ballov['Пуд'], 2)
Raschet_ballov['Общий балл'] = round((Raschet_ballov['К1']+Raschet_ballov['К2']+Raschet_ballov['К3']+Raschet_ballov['К4']+Raschet_ballov['К5'])/5, 2)



Answers_respond_list_st = Answers_respond_st.columns.tolist() ##извлекаем наименования столбцов в список

New_col = []  # Создаем пустой список
for i in range(25):  # Цикл от 0 до 18
    sim = i   # присваиваем номер
    New_col.append('v' + str(sim))  # добавляем новый номер вопрса в список

dictionary = dict(zip(Answers_respond_list_st, New_col)) # создаем  словарь для переименования стобцов
Answers_respond_st = Answers_respond_st.rename(columns=dictionary) # переименовываем столбцы в начальном датафрейме
Answers_respond_st['v25'] = Answers_respond_st['v2'].map({'менее 7 календарных дней': 6, '7 календарных дней': 7, '10 календарных дней': 10, '12 календарных дней': 12, '13 календарных дней': 13, '14 календарных дней и более': 14})

# Рассчитываем значение для нового столбца
def calculate_value1(row):
    total_answers = sum([val for val in row['v25'] if not math.isnan(val)])

    result_sum = sum([(i+1)*val for i, val in enumerate(row['v25']) if not math.isnan(val)])
    result = round(result_sum / total_answers)
    
    if result == 14:
        return 10
    elif result == 13:
        return 20
    elif result == 12:
        return 40
    elif 11 <= result <= 8:
        return 60
    else:
        return 100


result_df1 = Answers_respond_st.groupby('v0').apply(calculate_value1).reset_index()
result_df1.columns = ['v0', 'v25']

# Создание нового DataFrame для хранения результатов подсчета, считам количество ответов да на вопросы анкеты
ans_res_st = pd.DataFrame({'v0': Answers_respond_st['v0'].unique()})

selected_columns = ['v1', 'v3', 'v5', 'v7', 'v8', 'v12', 'v13', 'v14', 'v15', 'v16', 'v17', 'v19', 'v20', 'v21', 'v22', 'v23', 'v24']
# Используем цикл для подсчета значений и создания новых столбцов
for col in selected_columns:
    value = 'да'  # Значение, которое мы считаем
    count_col_name = f'_{col}_'
    counts = Answers_respond_st[Answers_respond_st[col] == value].groupby('v0').size().reset_index(name=count_col_name)
    ans_res_st = ans_res_st.merge(counts, on='v0', how='left')

ans_res_st = ans_res_st.merge(result_df1, on='v0', how='left')

ans_res_st = ans_res_st.dropna(axis=1) # Удаляем столбцы со значением NaN
ans_res_st['v0'] = ans_res_st['v0'].str.replace('.', '')# Удаляем точку из наименований организаций
ans_res_st = ans_res_st.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
ans_res_st = ans_res_st.reset_index(drop=True)

col_ob1 = Answers_respond_st.groupby('v0').size().reset_index(name='Ч_общ')
col_ob1['v0'] = col_ob1['v0'].str.replace('.', '')# Удаляем точку из наименований организаций
col_ob1 = col_ob1.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
col_ob1 = col_ob1.reset_index(drop=True)
all_ans1 = col_ob1['Ч_общ']

name_org = chek_list_st.filter(like='Наименование организации').copy()

Raschet_ballov1 = name_org
Raschet_ballov1['Истенд'] = chek_list_st.filter(like='На СТЕНДЕ').sum(axis=1)
Raschet_ballov1['Исайт'] = chek_list_st.filter(like='На САЙТЕ').sum(axis=1)
Raschet_ballov1['Инорм-стенд'] = chek_list_st.filter(like='На СТЕНДЕ').count(axis=1)
Raschet_ballov1['Инорм-сайт'] = chek_list_st.filter(like='На САЙТЕ').count(axis=1)
Raschet_ballov1['Пинф'] = round(0.5*((Raschet_ballov1['Истенд']/Raschet_ballov1['Инорм-стенд'])+(Raschet_ballov1['Исайт']/Raschet_ballov1['Инорм-сайт']))*100, 2)
Raschet_ballov1['Тдист'] = 30
Raschet_ballov1['Сдист'] = chek_list_st.filter(like='Наличие и функционирование на официальном сайте').sum(axis=1)
Raschet_ballov1['Пдист'] = Raschet_ballov1['Тдист']*Raschet_ballov1['Сдист']
Raschet_ballov1['Пдист'].where(Raschet_ballov1['Пдист'] <= 100, 100, inplace=True)
Raschet_ballov1['Устенд'] = ans_res_st['_v14_']
Raschet_ballov1['Усайт'] = ans_res_st['_v16_']
Raschet_ballov1['Уобщ-стенд'] = ans_res_st['_v13_']
Raschet_ballov1['Уобщ-сайт'] = ans_res_st['_v15_']
Raschet_ballov1['Поткруд'] = round(0.5*((Raschet_ballov1['Устенд']/Raschet_ballov1['Уобщ-стенд'])+(Raschet_ballov1['Усайт']/Raschet_ballov1['Уобщ-сайт']))*100, 2)
Raschet_ballov1['К1'] = round(0.3*Raschet_ballov1['Пинф'] + 0.3*Raschet_ballov1['Пдист'] + 0.4*Raschet_ballov1['Поткруд'], 2)
Raschet_ballov1['Ткомф'] = chek_list_st.filter(like='Обеспечение в организации комфортных условий').sum(axis=1)
Raschet_ballov1['Скомф'] = 20
Raschet_ballov1['Пкомф.усл'] = Raschet_ballov1['Ткомф']*Raschet_ballov1['Скомф']
Raschet_ballov1['Пкомф.усл'].where(Raschet_ballov1['Пкомф.усл'] <= 100, 100, inplace=True)
Raschet_ballov1['ожид'] = ans_res_st['v25']
Raschet_ballov1['Усвоевр'] = ans_res_st['_v3_']
Raschet_ballov1['Чобщ1'] = ans_res_st['_v1_']
Raschet_ballov1['Пожид'] = (round(Raschet_ballov1['Усвоевр']/Raschet_ballov1['Чобщ1']*100, 2) + Raschet_ballov1['ожид'])/2
Raschet_ballov1['Укомф'] = (ans_res_st['_v5_']+ans_res_st['_v17_'])/2
Raschet_ballov1['Чобщ2'] = all_ans1
Raschet_ballov1['Пкомфуд'] = round(Raschet_ballov1['Укомф']/Raschet_ballov1['Чобщ2']*100, 2)
Raschet_ballov1['К2'] = round(0.3*Raschet_ballov1['Пкомф.усл'] + 0.4*Raschet_ballov1['Пожид'] + 0.3*Raschet_ballov1['Пкомфуд'], 2)
Raschet_ballov1['Торгдост'] = chek_list_st.filter(like='Оборудование территории').sum(axis=1)
Raschet_ballov1['Соргдост'] = 20
Raschet_ballov1['Поргдост'] = Raschet_ballov1['Торгдост']*Raschet_ballov1['Соргдост']
Raschet_ballov1['Поргдост'].where(Raschet_ballov1['Поргдост'] <= 100, 100, inplace=True)
Raschet_ballov1['Туслугдост'] = chek_list_st.filter(like='Обеспечение в организации условий доступности').sum(axis=1)
Raschet_ballov1['Суслугдост'] = 20
Raschet_ballov1['Пуслугдост'] = Raschet_ballov1['Туслугдост']*Raschet_ballov1['Суслугдост']
Raschet_ballov1['Пуслугдост'].where(Raschet_ballov1['Пуслугдост'] <= 100, 100, inplace=True)
Raschet_ballov1['Удост'] = ans_res_st['_v12_']
Raschet_ballov1['Чинв'] = ans_res_st['_v8_']
Raschet_ballov1['Пдостуд'] = round(Raschet_ballov1['Удост']/Raschet_ballov1['Чинв']*100, 2)
Raschet_ballov1['К3'] = round(0.3*Raschet_ballov1['Поргдост'] + 0.4*Raschet_ballov1['Пуслугдост'] + 0.3*Raschet_ballov1['Пдостуд'], 2)
Raschet_ballov1['Уперв.конт'] = ans_res_st['_v7_'] 
Raschet_ballov1['Чобщ3'] = all_ans1
Raschet_ballov1['Пперв.контуд'] = round(Raschet_ballov1['Уперв.конт']/Raschet_ballov1['Чобщ3']*100, 2)
Raschet_ballov1['Уоказ.услуг'] = ans_res_st['_v19_']
Raschet_ballov1['Чобщ4'] = all_ans1
Raschet_ballov1['Показ.услугуд'] = round(Raschet_ballov1['Уоказ.услуг']/Raschet_ballov1['Чобщ4']*100, 2)
Raschet_ballov1['Увежл.дист'] = ans_res_st['_v24_']
Raschet_ballov1['Чобщ_ус'] = ans_res_st['_v23_']
Raschet_ballov1['Пвежл.дистуд'] = round(Raschet_ballov1['Увежл.дист']/Raschet_ballov1['Чобщ_ус']*100, 2)
Raschet_ballov1['К4'] = round(0.4*Raschet_ballov1['Пперв.контуд'] + 0.4*Raschet_ballov1['Показ.услугуд'] + 0.2*Raschet_ballov1['Пвежл.дистуд'], 2)
Raschet_ballov1['Уреком'] = ans_res_st['_v20_']
Raschet_ballov1['Чобщ5'] = all_ans1
Raschet_ballov1['Преком'] = round(Raschet_ballov1['Уреком']/Raschet_ballov1['Чобщ5']*100, 2)
Raschet_ballov1['Уорг.усл'] = ans_res_st['_v21_']
Raschet_ballov1['Чобщ6'] = all_ans1
Raschet_ballov1['Порг.услуд'] = round(Raschet_ballov1['Уорг.усл']/Raschet_ballov1['Чобщ6']*100, 2)
Raschet_ballov1['Ууд'] = ans_res_st['_v22_']
Raschet_ballov1['Чобщ7'] = all_ans1
Raschet_ballov1['Пуд'] = round(Raschet_ballov1['Ууд']/Raschet_ballov1['Чобщ7']*100, 2)
Raschet_ballov1['К5'] = round(0.3*Raschet_ballov1['Преком'] + 0.2*Raschet_ballov1['Порг.услуд'] + 0.5*Raschet_ballov1['Пуд'], 2)
Raschet_ballov1['Общий балл'] = round((Raschet_ballov1['К1']+Raschet_ballov1['К2']+Raschet_ballov1['К3']+Raschet_ballov1['К4']+Raschet_ballov1['К5'])/5, 2)

# Создание копии Raschet_ballov
Raschet_ballov2 = Raschet_ballov.copy()

# Соединение датафреймов и выполнение операции
for col in Raschet_ballov2.columns[1:]:
    if col in Raschet_ballov1.columns:
        Raschet_ballov2[col] = (Raschet_ballov[col].astype(float) + Raschet_ballov1[col].astype(float)) / 2
    else:
        Raschet_ballov2[col] = Raschet_ballov[col]
Raschet_ballov = Raschet_ballov2

row_chek_list = chek_list.columns.tolist()

New_col_for_chek_list = []  # Создаем пустой список
for i in range(chek_list.shape[1]+1):  # Цикл от 0 до 18
    sim = i   # присваиваем номер
    New_col_for_chek_list.append('us' + str(sim))  # добавляем новый номер вопрса в список

dict_chek = dict(zip(row_chek_list, New_col_for_chek_list))
chek_list = chek_list.rename(columns=dict_chek) # переименовываем столбцы в начальном датафрейме


name_org1 = pd.DataFrame({'us0': chek_list['us0']}) 
chek_list_stend = chek_list.iloc[:, 1:27]  # Датафрейм с 1-5 столбцами
chek_list_sait = chek_list.iloc[:, 28:56]  # Датафрейм с 6-10 столбцами = df.iloc[:, 0:5]  # Датафрейм с 1-5 столбцами
chek_list_dist = chek_list.iloc[:, 57:62]  # Датафрейм с 6-10 столбцами
chek_list_komf = chek_list.iloc[:, 63:69]
chek_list_obor_inv = chek_list.iloc[:, 70:74]
chek_list_sreda_inv = chek_list.iloc[:, 75:80]

chek_list_stend = pd.concat([name_org1, chek_list_stend], axis=1)
chek_list_sait = pd.concat([name_org1, chek_list_sait], axis=1)
chek_list_dist = pd.concat([name_org1, chek_list_dist], axis=1)
chek_list_komf = pd.concat([name_org1, chek_list_komf], axis=1)
chek_list_obor_inv = pd.concat([name_org1, chek_list_obor_inv], axis=1)
chek_list_sreda_inv = pd.concat([name_org1, chek_list_sreda_inv], axis=1)

def process_row(row):
    first_column_value = row.iloc[0]
    row_subset = row.iloc[1:]

    zero_values = row_subset[row_subset == 0].index.tolist()  # Получаем список индексов столбцов с нулевыми значениями

    return pd.Series([first_column_value, zero_values], index=['First_Column_Value', 'Zero_Values'])
# Применяем функцию и конвертируем результаты в DataFrame
nedostatki_stend = chek_list_stend.apply(process_row, axis=1)
nedostatki_stend = nedostatki_stend[nedostatki_stend.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_stend = nedostatki_stend.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))

nedostatki = nedostatki_stend
nedostatki = nedostatki.rename(columns={'First_Column_Value':'Name_org', 'Zero_Values':'bad_stend'})

nedostatki_sait = chek_list_sait.apply(process_row, axis=1)
nedostatki_sait = nedostatki_sait[nedostatki_sait.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_sait = nedostatki_sait.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_sait'] = nedostatki_sait['Zero_Values']

nedostatki_dist = chek_list_dist.apply(process_row, axis=1)
nedostatki_dist = nedostatki_dist[nedostatki_dist.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_dist = nedostatki_dist.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_dist'] = nedostatki_dist['Zero_Values']

nedostatki_komf = chek_list_komf.apply(process_row, axis=1)
nedostatki_komf = nedostatki_komf[nedostatki_dist.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_komf = nedostatki_komf.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_komf'] = nedostatki_komf['Zero_Values']

nedostatki_obor_inv = chek_list_obor_inv.apply(process_row, axis=1)
nedostatki_obor_inv = nedostatki_obor_inv[nedostatki_obor_inv.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_obor_inv = nedostatki_obor_inv.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['obor_inv'] = nedostatki_obor_inv['Zero_Values']

nedostatki_sreda_inv = chek_list_sreda_inv.apply(process_row, axis=1)
nedostatki_sreda_inv = nedostatki_sreda_inv[nedostatki_sreda_inv.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_sreda_inv = nedostatki_sreda_inv.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['sreda_inv'] = nedostatki_sreda_inv['Zero_Values']

row_chek_list1 = row_chek_list.copy
new_list = [re.search(r'\[([^\]]+)\]', item).group(1) for item in row_chek_list[1:] if re.search(r'\[([^\]]+)\]', item)]
name_list = [row_chek_list[0]]
row_chek_list = name_list + new_list
dict_chek1 = dict(zip(New_col_for_chek_list, row_chek_list))

output_data = []

for index, row in nedostatki.iterrows():
    output_row = [row['Name_org']]
    for col in nedostatki.columns[1:]:
        if row[col] == "нет недостатков":
            output_row.append(row[col])
        else:
            values_to_find = row[col] 
            result_keys = str([dict_chek1[value] for value in values_to_find ])
            output_row.append(result_keys)
    output_data.append(output_row)

output_df = pd.DataFrame(output_data, columns=['Name_org', 'bad_stend', 'bad_sait', 'bad_dist', 'bad_komf', 'obor_inv', 'sreda_inv'])

otchet = Document()
section = otchet.sections[0]
section.left_margin = Cm(3.0)  # Левое поле 2 дюйма
section.right_margin = Cm(1.5)  # Правое поле 2 дюйма
section.top_margin = Cm(2.0)  # Верхнее поле 2 дюйма
section.bottom_margin = Cm(2.0)  # Нижнее поле 2 дюйма
# Устанавливаем шрифт по умолчанию для дальнейшего добавляемого текста
default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

# Добавляем параграф с измененными параметрами шрифта
zag = otchet.add_paragraph()
run = zag.add_run("Основные результаты исследования")
font = run.font
run.bold = True
font.size = Pt(18) 
zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Добавляем параграф с измененными параметрами шрифта
under_zag = otchet.add_paragraph()
run = under_zag.add_run("Результаты независимой оценки качества условий оказания услуг медицинскими учреждениями")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 1. Открытость и доступность информации об медицинском учреждении")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

# Добавляем текст с установленным шрифтом
abz =otchet.add_paragraph("Критерий представлен тремя показателями:")
abz1 = otchet.add_paragraph("Показатель 1.1.	 Соответствие информации о деятельности медицинского учреждения, размещенной на общедоступных информационных ресурсах, ее содержанию и порядку (форме), установленным нормативными правовыми актами (на информационных стендах в помещении медицинского учреждения; на официальном сайте медицинского учреждения в сети «Интернет»).")
abz2 = otchet.add_paragraph("Показатель 1.2. 	Наличие на официальном сайте медицинского учреждения информации о дистанционных способах обратной связи и взаимодействия с получателями услуг и их функционирование (абонентского номера телефона; адреса электронной почты; электронных сервисов (для подачи электронного обращения (жалобы, предложения), получения консультации по оказываемым услугам и иных.); раздела официального сайта «Часто задаваемые вопросы»; технической возможности выражения получателем услуг мнения о качестве условий оказания услуг медицинским учреждением (наличие анкеты для опроса граждан или гиперссылки на нее)).")
abz3 = otchet.add_paragraph("Показатель 1.3.	 Доля получателей услуг, удовлетворенных открытостью, полнотой и доступностью информации о деятельности медицинского учреждения, размещенной на информационных стендах в помещении медицинского учреждения, на официальном сайте медицинского учреждения в сети «Интернет» (в % от общего числа опрошенных получателей услуг).")
abz4 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz5 = otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пинф', 'Пдист', 'Поткруд', 'К1']]
min_value = table['К1'].min()
max_value = table['К1'].max()
mean_value = table['К1'].mean()
sorted_table = table.sort_values(by='К1', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz6 = otchet.add_paragraph(f"Итоговые баллы по критерию «Открытость и доступность информации о медицинском учреждении» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz7 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К1']}балла.")    
abz8 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
   otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К1']}балла.")

# Добавляем таблицу в документ
table1 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table1.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table1.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table1.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 2. Комфортность условий предоставления услуг")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz10 = otchet.add_paragraph("Критерий представлен двумя показателями:")
abz11 = otchet.add_paragraph("Показатель 2.1. Обеспечение в медицинском учреждении комфортных условий пребывания (транспортная/ пешая доступность медицинского учреждения, санитарное состояние помещений и территории учреждения, наличие и доступность питьевой воды, санитарно-гигиенических помещений, достаточность гардеробов)")
abz12 = otchet.add_paragraph("Показатель 2.3. Доля получателей услуг, удовлетворенных комфортностью предоставления услуг медицинским учреждением (в % от общего числа опрошенных получателей услуг).")
abz13 = otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table11 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пкомф.усл', 'Пожид', 'Пкомфуд', 'К2']]
min_value = table11['К2'].min()
max_value = table11['К2'].max()
mean_value = table11['К2'].mean()
sorted_table = table11.sort_values(by='К2', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz13_1 = otchet.add_paragraph(f"Итоговые баллы по критерию «Комфортность условий предоставления услуг» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz13_2 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К2']}балла.")    
abz13_3 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
   otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К2']}балла.")
    
# Добавляем таблицу в документ
table2 =otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table2.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table2.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table2.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 3. Доступность услуг для инвалидов")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz20 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz21 = otchet.add_paragraph("Показатель 3.1. Оборудование помещений медицинского учреждения и прилегающей к ней территории с учетом доступности для инвалидов (наличие оборудованных входных групп пандусами (подъемными платформами); наличие выделенных стоянок для автотранспортных средств инвалидов; наличие адаптированных лифтов, поручней, расширенных дверных проемов; наличие сменных кресел-колясок; наличие специально оборудованных санитарно-гигиенических помещений в учреждения социальной сферы).")
abz22 = otchet.add_paragraph("Показатель 3.2. Обеспечение в медицинском учреждении условий доступности, позволяющих инвалидам получать услуги наравне с другими (дублирование для инвалидов по слуху и зрению звуковой и зрительной информации; дублирование надписей, знаков и иной текстовой и графической информации знаками, выполненными рельефно-точечным шрифтом Брайля; возможность предоставления инвалидам по слуху (слуху и зрению) услуг сурдопереводчика (тифлосурдопереводчика); наличие альтернативной версии официального сайта учреждения социальной сферы в сети «Интернет» для инвалидов по зрению; помощь, оказываемая работниками учреждения социальной сферы, прошедшими необходимое обучение (инструктирование) по сопровождению инвалидов в помещениях учреждения социальной сферы и на прилегающей территории; наличие возможности предоставления услуги в дистанционном режиме или на дому). ")
abz23 = otchet.add_paragraph("Показатель 3.3. Доля получателей услуг, удовлетворенных доступностью услуг для инвалидов (в % от общего числа опрошенных получателей услуг – инвалидов)")
abz24 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table12 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Поргдост', 'Пуслугдост', 'Пдостуд', 'К3']]
min_value = table12['К3'].min()
max_value = table12['К3'].max()
mean_value = table12['К3'].mean()
sorted_table = table12.sort_values(by='К3', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz25 = otchet.add_paragraph(f"Итоговые баллы по критерию «Доступность услуг для инвалидов» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz26 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К3']}балла.") 
abz27 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К3']}балла.")

table3 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table3.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table3.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table3.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 4. Доброжелательность, вежливость работников учреждения")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz30 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz31 = otchet.add_paragraph("Показатель 4.1. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников медицинского учреждения, обеспечивающих первичный контакт и информирование получателя услуги при непосредственном обращении в организацию образования (в % от общего числа опрошенных получателей услуг)")
abz32 = otchet.add_paragraph("Показатель 4.2. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников медицинского учреждения, обеспечивающих непосредственное оказание услуги при обращении в организацию образования (в % от общего числа опрошенных получателей услуг)")
abz33 = otchet.add_paragraph("Показатель 4.3. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников медицинского учреждения при использовании дистанционных форм взаимодействия (в % от общего числа опрошенных получателей услуг).")
abz34 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table13 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пперв.контуд', 'Показ.услугуд', 'Пвежл.дистуд', 'К4']]
min_value = table13['К4'].min()
max_value = table13['К4'].max()
mean_value = table13['К4'].mean()
sorted_table = table13.sort_values(by='К4', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz35 = otchet.add_paragraph(f"Итоговые баллы по критерию «Доброжелательность, вежливость работников учреждения» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz36 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К4']}балла.") 
abz37 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К4']}балла.")

table4 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table4.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table4.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table4.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 5. Удовлетворенность условиями оказания услуг")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz40 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz41 = otchet.add_paragraph("Показатель 5.1. Доля получателей услуг, которые готовы рекомендовать социальное учреждение родственникам и знакомым")
abz42 = otchet.add_paragraph("Показатель 5.2. Доля получателей услуг, удовлетворенных организационными условиями предоставления услуг (графиком и режимом работы социального учреждения) (в % от общего числа опрошенных получателей услуг)")
abz43 = otchet.add_paragraph("Показатель 5.3. Доля получателей услуг, удовлетворенных в целом условиями оказания услуг в медицинском учреждении (в % от общего числа опрошенных получателей услуг).")
abz44 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table14 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Преком', 'Порг.услуд', 'Пуд', 'К5']]
min_value = table14['К5'].min()
max_value = table14['К5'].max()
mean_value = table14['К5'].mean()
sorted_table = table14.sort_values(by='К5', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz45 = otchet.add_paragraph(f"Итоговые баллы по критерию «Удовлетворенность условиями оказания услуг» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz46 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К5']}балла.") 
abz47 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К5']}балла.")

table5 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table5.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table5.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table5.cell(i+1, j).text = str(sorted_table.values[i, j])

under_zag = otchet.add_paragraph()
run = under_zag.add_run("Итоговая оценка качества условий оказания услуг медицинскими учреждениями. Рейтинг учреждений")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Общий рейтинг медицинских учреждений.")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

table15 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'К1', 'К2', 'К3', 'К4', 'К5', 'Общий балл']]
min_value = table15['Общий балл'].min()
max_value = table15['Общий балл'].max()
mean_value = table15['Общий балл'].mean()
sorted_table = table15.sort_values(by='Общий балл', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(1)

sorted_table['рейтинг'] = range(1, len(sorted_table) + 1)

table6 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table6.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table6.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table6.cell(i+1, j).text = str(sorted_table.values[i, j])

abz45 = otchet.add_paragraph(f"Итоговый анализ и оценка качества работы социальных учреждений позволяет определить лучшие учреждения по результатам мониторинга.  Общий балл организаций варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по сумме критериев {mean_value}.")
abz46 = otchet.add_paragraph(f"Среди социальных учреждений {plase} в первую тройку лидеров вошли следующие учреждения:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['Общий балл']}балла.") 
abz47 = otchet.add_paragraph("Последнюю строку рейтинга занимает")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['Общий балл']}балла.")

under_zag = otchet.add_paragraph()
run = under_zag.add_run(f"Основные выводы и рекомендации по результатам независимой оценки качества условий оказания услуг медицинских учреждений {plase}.")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
under_zag = otchet.add_paragraph()
run = under_zag.add_run(f"Основные выводы по результатам независимой оценки.")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

list_krit = {'Наименование критерия':['Критерий 1', 'Критерий 2', 'Критерий 3', 'Критерий 4', 'Критерий 5'],
             'Значение':[table['К1'].mean(), table11['К2'].mean(), table12['К3'].mean(), table13['К4'].mean(), table14['К5'].mean()]}
list_krit = pd.DataFrame(list_krit)
sorted_list_krit = list_krit.sort_values(by='Значение')
sorted_list_krit = sorted_list_krit.reset_index(drop=True)

# Создаем словарь для массовой замены значений
list_krit_dict = {'Критерий 1': '1. Открытость и доступность информации', 
                  'Критерий 2': '2. Комфортность условий предоставления услуг', 
                  'Критерий 3': '3. Доступность услуг для инвалидов', 
                  'Критерий 4': '4. Доброжелательность, вежливость работников организации', 
                  'Критерий 5': '5. Удовлетворенность условиями оказания услуг'
                 }
# Массово заменяем значения в столбце 'Имя'
sorted_list_krit['Наименование критерия'] = sorted_list_krit['Наименование критерия'].replace(list_krit_dict)

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz50 = otchet.add_paragraph(f"Согласно результатам проведённого исследования, основным недостатком у данных учреждений является {sorted_list_krit.at[0, 'Наименование критерия']}. ")
abz51 = otchet.add_paragraph(f"Также есть проблемы с {sorted_list_krit.at[1, 'Наименование критерия']} и {sorted_list_krit.at[2, 'Наименование критерия']}. ")# Вставка графика в документ Word

if not Raschet_ballov.empty:
    print("DataFrame не является пустым")
else:
    print("DataFrame пустой")

if not sorted_table.empty:
    print("DataFrame не является пустым")
else:
    print("DataFrame пустой")

# Создаем таблицу с нужным количеством строк и столбцов
table20 = otchet.add_table(rows=1, cols=4)
table20.style = 'Table Grid'  # Применяем стиль таблицы

# Заголовки столбцов
hdr_cells = table20.rows[0].cells
hdr_cells[0].text = 'Наименование организации'
hdr_cells[1].text = 'Балл'
hdr_cells[2].text = 'Рейтинг'
hdr_cells[3].text = 'Недостатки'

# Проходимся по каждой строке и добавляем данные в таблицу
for index, row in output_df.iterrows():
    row_cells = table20.add_row().cells
    row_cells[0].text = str(row['Name_org'])
    row_cells[1].text = str(Raschet_ballov.loc[Raschet_ballov['Наименование организации или П/Н по списку'] == row['Name_org'], 'Общий балл'].values[0])
    row_cells[2].text = str(sorted_table.loc[sorted_table['Наименование организации или П/Н по списку'] == row['Name_org'], 'рейтинг'].values[0])

    row_cells[3].text = f"Недостатки на стенде: отсутствуют документы о {str(row['bad_stend'])}\n"\
                        f"Недостатки на сайте: отсутствуют документы о {str(row['bad_sait'])}\n"\
                        f"Недостатки функционирование дистанционных способов связи: {str(row['bad_dist'])}\n"\
                        f"Недостатки комфортности условий предоставления услуг: {str(row['bad_komf'])}\n"\
                        f"Недостатки в разрезе оборудования для инвалидов: {str(row['obor_inv'])}\n"\
                        f"Недостатки доступности среды для инвалидов: {str(row['sreda_inv'])}"



# Создаем таблицу с нужным количеством строк и столбцов
table30 = otchet.add_table(rows=1, cols=2)
table30.style = 'Table Grid'  # Применяем стиль таблицы

# Заголовки столбцов
hdr_cells = table30.rows[0].cells
hdr_cells[0].text = 'Наименование организации'
hdr_cells[1].text = 'Рекомендации'

# Проходимся по каждой строке и добавляем данные в таблицу
for index, row in output_df.iterrows():
    row_cells = table30.add_row().cells
    row_cells[0].text = str(row['Name_org'])
    
    K3_list = table15.loc[table15['Наименование организации или П/Н по списку'] == row['Name_org'], 'К3'].values[0]
        # Добавляем результат проверки условия в столбец "Недостатки"
    if K3_list < 90:
        result_k3 = ''.join(("Провести инструктаж и обучение сотрудников по вопросам обеспечения доступности организации для инвалидов, в т.ч. ",
                             "по сопровождению маломобильных граждан в помещениях организации и на прилегающей территории с учетом имеющихся у ",
                             "них ограничений и расстройств, по оказанию содействия инвалиду в беспрепятственном получении услуги.",
                             "определить и закрепить в должностных инструкциях и распорядительных документах ответственность конкретных работников ",
                             "(категорий работников) организации по сопровождению маломобильных граждан и оказанию им содействия в получении услуги.",
                             "Провести оптимизацию расположения кабинетов организации с учетом обеспечения их доступности для маломобильных граждан, ",
                             "перенести наиболее востребованные кабинеты и услуги на первый этаж"))
    else:
        result_k3 = 'нет рекомендаций'
    # Добавляем результат проверки условия в столбец "Недостатки"

    K4_list = table15.loc[table15['Наименование организации или П/Н по списку'] == row['Name_org'], 'К4'].values[0]
    
    if K4_list < 90:
        result_k4 = ''.join(("Провести обучение персонала организации по вопросам этики и деонтологии",
                             "Провести обучение персонала организации по вопросам клиентоцентричности, реализации стандарта «Государство для людей» ",
                             "и принципов проактивного предоставления услуг. Ввести на регулярной основе (не реже 1 раза в квартал) рабочие совещания ",
                             "с коллективом организации по вопросам соблюдения норм профессиональной этики и правил служебного поведения",
                             "Разработка внутреннего стандарта и памятки по предоставлению информации по телефону для специалистов организации",
                             "Внести изменения в трудовой договор с работниками организации, включив результаты НОК (отдельные критерии/показатели) ",
                             "в состав показателей результативности каждого работника, осуществляющего взаимодействие с получателями услуг, ",
                             "для установления стимулирующих выплат и использования в рамках программ нематериальной мотивации"))
    else:
        result_k4 = 'нет рекомендаций'    
    # Добавляем результат проверки условия в столбец "Недостатки"
    K5_list = table15.loc[table15['Наименование организации или П/Н по списку'] == row['Name_org'], 'К5'].values[0]
    
    if K5_list < 90:
        result_k5 = ''.join(("Провести внутренний аудит системы менеджмента качества в структурных подразделениях организации ",
                             "с целью реализации превентивных мер, направленных на совершенствование условий оказания услуг по перечню недостатков, ",
                             "выявленных в других организациях или структурных подразделениях (управление по прецедентам)",
                             "Провести анкетирование среди получателей социальных услуг, направленное на выявление глубинных причин ",
                             "неудовлетворенности условиями оказания услуг. Разработать карты клиентского пути, направленные на уточнение ",
                             "и оптимизацию сценариев и условий оказания услуг в организации. Разработать и провести информационную кампанию ",
                             "по информированию граждан о формах и видах оказываемых услуг, существующих преимуществах получения их в данной организации, ",
                             "возможных льготах и пр. (изготовление памяток, буклетов, написание статей в СМИ и т.п.) ",
                             "Внедрить в практику работы организации проекты, направленные на повышение эффективности предоставления социальных услуг ",
                             "и работу с целевыми категориями получателей услуг (Например, созданию условий для обучения людей с ментальными нарушениями, ",
                             "Территория возможностей для граждан серебряного возраста по организации и поддержке творческого досуга граждан пожилого возраста, ",
                             "удовлетворению и развитию их культурных, духовных потребностей). На уровне уполномоченного органа власти (учредителя организации) ",
                             "внести изменения в трудовой договор с руководителями организаций социальной сферы, включив общие результаты НОК ",
                             "и результаты исполнения планов по устранению недостатков, выявленных в ходе НОК в показатели результативности для установления ",
                             "стимулирующих выплат и использования в рамках программ нематериальной мотивации"))
    else:
        result_k5 = 'нет рекомендаций'
        
    row_cells[1].text = f"Разместить на стенде организации следующие документы: {str(row['bad_stend'])}\n"\
                        f"Разместить на сайте организации следующие документы: {str(row['bad_sait'])}\n"\
                        f"Обеспечить наличие в организации следующих видов дистанционного взаимодействия: {str(row['bad_dist'])}\n"\
                        f"В разрезе удовлетворенности доступностью услуг для инвалидов рекомендовано: {result_k3}\n"\
                        f"В разрезе удовлетворенности вежливостью и доброжелательностью работников учреждения рекомендовано:: {result_k4}\n"\
                        f"В разрезе удовлетворенности предоставлением услуг в целом рекомендовано: {result_k5}"



button = st.button("получить готовый файл расчет баллов")

if button:
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    Raschet_ballov.to_excel(writer, index=False, sheet_name='Sheet1')
    writer.close()
    output.seek(0)
    st.download_button(
        label="Загрузить",
        data=output,
        file_name='результаты.xlsx',
        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


button = st.button("Получить готовый отчет")
if button:
    bio = io.BytesIO()
    otchet.save(bio)
    st.download_button(
        label="Скачать",
        data=bio.getvalue(),
        file_name="Отчет.docx",
        mime="docx"
    )

