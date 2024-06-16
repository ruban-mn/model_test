import pandas as pd
import numpy as np
import openpyxl
import docx
import streamlit as st
import io
import time as tm
import re
from docx import Document
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# In[7]:


st.subheader('Независимая оценка качества услуг образовательных организаций')


# In[6]:

uploaded_file = st.file_uploader("**Загрузите файл сводную по чек-листам**", type=["xls", "xlsx"])


if uploaded_file is not None:
    # Чтение данных из загруженного файла Excel
    chek_list = pd.read_excel(uploaded_file)


# In[7]:


uploaded_file1 = st.file_uploader("**Загрузите файл с массивом анкет**", type=["xls", "xlsx"])

if uploaded_file1 is not None:
    # Чтение данных из загруженного файла Excel
    Answers_respond = pd.read_excel(uploaded_file1)


# In[8]:
plase = st.text_input('Введите название территории в родительном падеже', 'н-р Московской области')

## подгружаем сводную по чек листам
##chek_list=pd.read_excel(r"C:\Users\user\чек лист гулькевичи.xlsx")
##подгружаем массив с ответами респондентов
##Answers_respond=pd.read_excel(r"C:\Users\user\Анкета Гулькевичи НОК культура (Ответы).xlsx")


# In[ ]:


tm.sleep(30)


# In[9]:


Answers_respond_list = Answers_respond.columns.tolist() ##извлекаем наименования столбцов в список


# In[ ]:


New_col = []  # Создаем пустой список
for i in range(18):  # Цикл от 0 до 18
    sim = i   # присваиваем номер
    New_col.append('v' + str(sim))  # добавляем новый номер вопрса в список


# In[ ]:


dictionary = dict(zip(Answers_respond_list, New_col)) # создаем  словарь для переименования стобцов


# In[ ]:


Answers_respond = Answers_respond.rename(columns=dictionary) # переименовываем столбцы в начальном датафрейме


# In[ ]:


# Создание нового DataFrame для хранения результатов подсчета, считам количество ответов да на вопросы анкеты
ans_res = pd.DataFrame({'v0': Answers_respond['v0'].unique()})

# Используем цикл для подсчета значений и создания новых столбцов
for col in New_col:
    value = 'Да'  # Значение, которое мы считаем
    count_col_name = f'_{col}_'
    counts = Answers_respond[Answers_respond[col] == value].groupby('v0').size().reset_index(name=count_col_name)
    ans_res = ans_res.merge(counts, on='v0', how='left')


# In[ ]:


ans_res = ans_res.dropna(axis=1) # Удаляем столбцы со значением NaN
ans_res['v0'] = ans_res['v0'].str.replace('.', '')# Удаляем точку из наименований организаций
ans_res = ans_res.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
ans_res = ans_res.reset_index(drop=True)


# In[ ]:


col_ob = Answers_respond.groupby('v0').size().reset_index(name='Ч_общ')
col_ob['v0'] = col_ob['v0'].str.replace('.', '')# Удаляем точку из наименований организаций
col_ob = col_ob.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
col_ob = col_ob.reset_index(drop=True)
all_ans = col_ob['Ч_общ']


# In[ ]:


name_org = chek_list.filter(like='Наименование организации').copy()


# In[10]:


Raschet_ballov = name_org
Raschet_ballov['Истенд'] = chek_list.filter(like='На СТЕНДЕ').sum(axis=1)
Raschet_ballov['Исайт'] = chek_list.filter(like='На САЙТЕ').sum(axis=1)
Raschet_ballov['Инорм-стенд'] = chek_list.filter(like='На СТЕНДЕ').count(axis=1)
Raschet_ballov['Инорм-сайт'] = chek_list.filter(like='На САЙТЕ').count(axis=1)
Raschet_ballov['Пинф'] = round(0.5*((Raschet_ballov['Истенд']/Raschet_ballov['Инорм-стенд'])+(Raschet_ballov['Исайт']/Raschet_ballov['Инорм-сайт']))*100, 2)
Raschet_ballov['Тдист'] = 30
Raschet_ballov['Сдист'] = chek_list.filter(like='Наличие и функционирование на официальном сайте').sum(axis=1)
Raschet_ballov['Пдист'] = Raschet_ballov['Тдист']*Raschet_ballov['Сдист']
Raschet_ballov['Пдист'].where(Raschet_ballov['Пдист'] <= 100, 100, inplace=True)
Raschet_ballov['Устенд'] = ans_res['_v2_']
Raschet_ballov['Усайт'] = ans_res['_v4_']
Raschet_ballov['Уобщ-стенд'] = ans_res['_v1_']
Raschet_ballov['Уобщ-сайт'] = ans_res['_v3_']
Raschet_ballov['Поткруд'] = round(0.5*((Raschet_ballov['Устенд']/Raschet_ballov['Уобщ-стенд'])+(Raschet_ballov['Усайт']/Raschet_ballov['Уобщ-сайт']))*100, 2)
Raschet_ballov['К1'] = round(0.3*Raschet_ballov['Пинф'] + 0.3*Raschet_ballov['Пдист'] + 0.4*Raschet_ballov['Поткруд'], 2)
Raschet_ballov['Ткомф'] = chek_list.filter(like='Обеспечение в организации комфортных условий').sum(axis=1)
Raschet_ballov['Скомф'] = 20
Raschet_ballov['Пкомф.усл'] = Raschet_ballov['Ткомф']*Raschet_ballov['Скомф']
Raschet_ballov['Пкомф.усл'].where(Raschet_ballov['Пкомф.усл'] <= 100, 100, inplace=True)
Raschet_ballov['Укомф'] = ans_res['_v5_']
Raschet_ballov['Чобщ'] = col_ob['Ч_общ']
Raschet_ballov['Пкомфуд'] = round(Raschet_ballov['Укомф']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['К2'] = round(0.5*Raschet_ballov['Пкомф.усл'] + 0.5*Raschet_ballov['Пкомфуд'], 2)
Raschet_ballov['Торгдост'] = chek_list.filter(like='Оборудование территории').sum(axis=1)
Raschet_ballov['Соргдост'] = 20
Raschet_ballov['Поргдост'] = Raschet_ballov['Торгдост']*Raschet_ballov['Соргдост']
Raschet_ballov['Туслугдост'] = chek_list.filter(like='Обеспечение в организации условий доступности').sum(axis=1)
Raschet_ballov['Суслугдост'] = 20
Raschet_ballov['Пуслугдост'] = Raschet_ballov['Туслугдост']*Raschet_ballov['Суслугдост']
Raschet_ballov['Пуслугдост'].where(Raschet_ballov['Пуслугдост'] <= 100, 100, inplace=True)
Raschet_ballov['Удост'] = ans_res['_v7_']
Raschet_ballov['Чинв'] = ans_res['_v6_']
Raschet_ballov['Пдостуд'] = round(Raschet_ballov['Удост']/Raschet_ballov['Чинв']*100, 2)
Raschet_ballov['К3'] = round(0.3*Raschet_ballov['Поргдост'] + 0.4*Raschet_ballov['Пуслугдост'] + 0.3*Raschet_ballov['Пдостуд'], 2)
Raschet_ballov['Уперв.конт'] = ans_res['_v8_']
Raschet_ballov['Чобщ1'] = all_ans
Raschet_ballov['Пперв.контуд'] = round(Raschet_ballov['Уперв.конт']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['Уоказ.услуг'] = ans_res['_v9_']
Raschet_ballov['Чобщ2'] = all_ans
Raschet_ballov['Показ.услугуд'] = round(Raschet_ballov['Уоказ.услуг']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['Увежл.дист'] = ans_res['_v11_']
Raschet_ballov['Чобщ_ус'] = ans_res['_v10_']
Raschet_ballov['Пвежл.дистуд'] = round(Raschet_ballov['Увежл.дист']/Raschet_ballov['Чобщ_ус']*100, 2)
Raschet_ballov['К4'] = round(0.4*Raschet_ballov['Пперв.контуд'] + 0.4*Raschet_ballov['Показ.услугуд'] + 0.2*Raschet_ballov['Пвежл.дистуд'], 2)
Raschet_ballov['Уреком'] = ans_res['_v12_']
Raschet_ballov['Чобщ3'] = all_ans
Raschet_ballov['Преком'] = round(Raschet_ballov['Уреком']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['Уорг.усл'] = ans_res['_v13_']
Raschet_ballov['Чобщ4'] = all_ans
Raschet_ballov['Порг.услуд'] = round(Raschet_ballov['Уорг.усл']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['Ууд'] = ans_res['_v14_']
Raschet_ballov['Чобщ5'] = all_ans
Raschet_ballov['Пуд'] = round(Raschet_ballov['Ууд']/Raschet_ballov['Чобщ']*100, 2)
Raschet_ballov['К5'] = round(0.3*Raschet_ballov['Преком'] + 0.2*Raschet_ballov['Порг.услуд'] + 0.5*Raschet_ballov['Пуд'], 2)
Raschet_ballov['Общий балл'] = round((Raschet_ballov['К1']+Raschet_ballov['К2']+Raschet_ballov['К3']+Raschet_ballov['К4']+Raschet_ballov['К5'])/5, 2)

row_chek_list = chek_list.columns.tolist()

New_col_for_chek_list = []  # Создаем пустой список
for i in range(chek_list.shape[1]+1):  # Цикл от 0 до 18
    sim = i   # присваиваем номер
    New_col_for_chek_list.append('us' + str(sim))  # добавляем новый номер вопрса в список

dict_chek = dict(zip(row_chek_list, New_col_for_chek_list))
chek_list = chek_list.rename(columns=dict_chek) # переименовываем столбцы в начальном датафрейме

name_org1 = pd.DataFrame({'us0': chek_list['us0']}) 
chek_list_stend = chek_list.iloc[:, 1:18]  # Датафрейм с 1-5 столбцами
chek_list_sait = chek_list.iloc[:, 18:37]  # Датафрейм с 6-10 столбцами = df.iloc[:, 0:5]  # Датафрейм с 1-5 столбцами
chek_list_dist = chek_list.iloc[:, 37:42]  # Датафрейм с 6-10 столбцами
chek_list_komf = chek_list.iloc[:, 42:48]
chek_list_obor_inv = chek_list.iloc[:, 48:53]
chek_list_sreda_inv = chek_list.iloc[:, 53:58]

chek_list_stend = pd.concat([name_org1, chek_list_stend], axis=1)
chek_list_sait = pd.concat([name_org1, chek_list_sait], axis=1)
chek_list_dist = pd.concat([name_org1, chek_list_dist], axis=1)
chek_list_komf = pd.concat([name_org1, chek_list_komf], axis=1)
chek_list_obor_inv = pd.concat([name_org1, chek_list_obor_inv], axis=1)
chek_list_sreda_inv = pd.concat([name_org1, chek_list_sreda_inv], axis=1)

def process_row(row):
    first_column_value = row.iloc[0]
    row_subset = row.iloc[1:]

    zero_values = row_subset[row_subset == 0].index.tolist()  # Получаем список индексов столбцов с нулевыми значениями

    return pd.Series([first_column_value, zero_values], index=['First_Column_Value', 'Zero_Values'])
# Применяем функцию и конвертируем результаты в DataFrame
nedostatki_stend = chek_list_stend.apply(process_row, axis=1)
nedostatki_stend = nedostatki_stend[nedostatki_stend.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_stend = nedostatki_stend.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))

nedostatki = nedostatki_stend
nedostatki = nedostatki.rename(columns={'First_Column_Value':'Name_org', 'Zero_Values':'bad_stend'})

nedostatki_sait = chek_list_sait.apply(process_row, axis=1)
nedostatki_sait = nedostatki_sait[nedostatki_sait.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_sait = nedostatki_sait.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_sait'] = nedostatki_sait['Zero_Values']

nedostatki_dist = chek_list_dist.apply(process_row, axis=1)
nedostatki_dist = nedostatki_dist[nedostatki_dist.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_dist = nedostatki_dist.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_dist'] = nedostatki_dist['Zero_Values']

nedostatki_komf = chek_list_komf.apply(process_row, axis=1)
nedostatki_komf = nedostatki_komf[nedostatki_dist.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_komf = nedostatki_komf.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_komf'] = nedostatki_komf['Zero_Values']

nedostatki_obor_inv = chek_list_obor_inv.apply(process_row, axis=1)
nedostatki_obor_inv = nedostatki_obor_inv[nedostatki_obor_inv.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_obor_inv = nedostatki_obor_inv.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['obor_inv'] = nedostatki_obor_inv['Zero_Values']

nedostatki_sreda_inv = chek_list_sreda_inv.apply(process_row, axis=1)
nedostatki_sreda_inv = nedostatki_sreda_inv[nedostatki_sreda_inv.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_sreda_inv = nedostatki_sreda_inv.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['sreda_inv'] = nedostatki_sreda_inv['Zero_Values']

row_chek_list1 = row_chek_list.copy
new_list = [re.search(r'\[([^\]]+)\]', item).group(1) for item in row_chek_list[1:] if re.search(r'\[([^\]]+)\]', item)]
name_list = [row_chek_list[0]]
row_chek_list = name_list + new_list
dict_chek1 = dict(zip(New_col_for_chek_list, row_chek_list))

output_data = []

for index, row in nedostatki.iterrows():
    output_row = [row['Name_org']]
    for col in nedostatki.columns[1:]:
        if row[col] == "нет недостатков":
            output_row.append(row[col])
        else:
            values_to_find = row[col] 
            result_keys = str([dict_chek1[value] for value in values_to_find ])
            output_row.append(result_keys)
    output_data.append(output_row)

output_df = pd.DataFrame(output_data, columns=['Name_org', 'bad_stend', 'bad_sait', 'bad_dist', 'bad_komf', 'obor_inv', 'sreda_inv'])

otchet = Document()
section = otchet.sections[0]
section.left_margin = Cm(3.0)  # Левое поле 2 дюйма
section.right_margin = Cm(1.5)  # Правое поле 2 дюйма
section.top_margin = Cm(2.0)  # Верхнее поле 2 дюйма
section.bottom_margin = Cm(2.0)  # Нижнее поле 2 дюйма
# Устанавливаем шрифт по умолчанию для дальнейшего добавляемого текста
default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

# Добавляем параграф с измененными параметрами шрифта
zag = otchet.add_paragraph()
run = zag.add_run("Основные результаты исследования")
font = run.font
run.bold = True
font.size = Pt(18) 
zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Добавляем параграф с измененными параметрами шрифта
under_zag = otchet.add_paragraph()
run = under_zag.add_run("Результаты независимой оценки качества условий оказания услуг образовательными учреждениями")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 1. Открытость и доступность информации об образовательном учреждении")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

# Добавляем текст с установленным шрифтом
abz =otchet.add_paragraph("Критерий представлен тремя показателями:")
abz1 = otchet.add_paragraph("Показатель 1.1.	 Соответствие информации о деятельности образовательного учреждения, размещенной на общедоступных информационных ресурсах, ее содержанию и порядку (форме), установленным нормативными правовыми актами (на информационных стендах в помещении учреждения образования; на официальном сайте образовательного учреждения в сети «Интернет»).")
abz2 = otchet.add_paragraph("Показатель 1.2. 	Наличие на официальном сайте образовательного учреждения информации о дистанционных способах обратной связи и взаимодействия с получателями услуг и их функционирование (абонентского номера телефона; адреса электронной почты; электронных сервисов (для подачи электронного обращения (жалобы, предложения), получения консультации по оказываемым услугам и иных.); раздела официального сайта «Часто задаваемые вопросы»; технической возможности выражения получателем услуг мнения о качестве условий оказания услуг учреждением культуры (наличие анкеты для опроса граждан или гиперссылки на нее)).")
abz3 = otchet.add_paragraph("Показатель 1.3.	 Доля получателей услуг, удовлетворенных открытостью, полнотой и доступностью информации о деятельности образовательного учреждения, размещенной на информационных стендах в помещении образовательного учреждения, на официальном сайте образовательного учреждения в сети «Интернет» (в % от общего числа опрошенных получателей услуг).")
abz4 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz5 = otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пинф', 'Пдист', 'Поткруд', 'К1']]
min_value = table['К1'].min()
max_value = table['К1'].max()
mean_value = table['К1'].mean()
sorted_table = table.sort_values(by='К1', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz6 = otchet.add_paragraph(f"Итоговые баллы по критерию «Открытость и доступность информации об образовательного учреждения» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz7 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К1']}балла.")    
abz8 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
   otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К1']}балла.")

# Добавляем таблицу в документ
table1 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table1.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table1.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table1.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 2. Комфортность условий предоставления услуг")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz10 = otchet.add_paragraph("Критерий представлен двумя показателями:")
abz11 = otchet.add_paragraph("Показатель 2.1. Обеспечение в образовательном учреждении комфортных условий пребывания (транспортная/ пешая доступность образовательного учреждения, санитарное состояние помещений и территории учреждения, наличие и доступность питьевой воды, санитарно-гигиенических помещений, достаточность гардеробов)")
abz12 = otchet.add_paragraph("Показатель 2.3. Доля получателей услуг, удовлетворенных комфортностью предоставления услуг образовательным учреждением (в % от общего числа опрошенных получателей услуг).")
abz13 = otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table11 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пкомф.усл', 'Пкомфуд', 'К2']]
min_value = table11['К2'].min()
max_value = table11['К2'].max()
mean_value = table11['К2'].mean()
sorted_table = table11.sort_values(by='К2', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

# Добавляем таблицу в документ
table2 =otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table2.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table2.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table2.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 3. Доступность услуг для инвалидов")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz20 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz21 = otchet.add_paragraph("Показатель 3.1. Оборудование помещений образовательного учреждения и прилегающей к ней территории с учетом доступности для инвалидов (наличие оборудованных входных групп пандусами (подъемными платформами); наличие выделенных стоянок для автотранспортных средств инвалидов; наличие адаптированных лифтов, поручней, расширенных дверных проемов; наличие сменных кресел-колясок; наличие специально оборудованных санитарно-гигиенических помещений в учреждения социальной сферы).")
abz22 = otchet.add_paragraph("Показатель 3.2. Обеспечение в учреждении образования условий доступности, позволяющих инвалидам получать услуги наравне с другими (дублирование для инвалидов по слуху и зрению звуковой и зрительной информации; дублирование надписей, знаков и иной текстовой и графической информации знаками, выполненными рельефно-точечным шрифтом Брайля; возможность предоставления инвалидам по слуху (слуху и зрению) услуг сурдопереводчика (тифлосурдопереводчика); наличие альтернативной версии официального сайта учреждения социальной сферы в сети «Интернет» для инвалидов по зрению; помощь, оказываемая работниками учреждения социальной сферы, прошедшими необходимое обучение (инструктирование) по сопровождению инвалидов в помещениях учреждения социальной сферы и на прилегающей территории; наличие возможности предоставления услуги в дистанционном режиме или на дому). ")
abz23 = otchet.add_paragraph("Показатель 3.3. Доля получателей услуг, удовлетворенных доступностью услуг для инвалидов (в % от общего числа опрошенных получателей услуг – инвалидов)")
abz24 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table12 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Поргдост', 'Пуслугдост', 'Пдостуд', 'К3']]
min_value = table12['К3'].min()
max_value = table12['К3'].max()
mean_value = table12['К3'].mean()
sorted_table = table12.sort_values(by='К3', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz25 = otchet.add_paragraph(f"Итоговые баллы по критерию «Доступность услуг для инвалидов» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz26 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К3']}балла.") 
abz27 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К3']}балла.")

table3 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table3.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table3.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table3.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 4. Доброжелательность, вежливость работников учреждения")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz30 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz31 = otchet.add_paragraph("Показатель 4.1. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников образовательного учреждения, обеспечивающих первичный контакт и информирование получателя услуги при непосредственном обращении в организацию образования (в % от общего числа опрошенных получателей услуг)")
abz32 = otchet.add_paragraph("Показатель 4.2. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников образовательного учреждения, обеспечивающих непосредственное оказание услуги при обращении в организацию образования (в % от общего числа опрошенных получателей услуг)")
abz33 = otchet.add_paragraph("Показатель 4.3. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников образовательного учреждения при использовании дистанционных форм взаимодействия (в % от общего числа опрошенных получателей услуг).")
abz34 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table13 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пперв.контуд', 'Показ.услугуд', 'Пвежл.дистуд', 'К4']]
min_value = table13['К4'].min()
max_value = table13['К4'].max()
mean_value = table13['К4'].mean()
sorted_table = table13.sort_values(by='К4', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz35 = otchet.add_paragraph(f"Итоговые баллы по критерию «Доброжелательность, вежливость работников учреждения» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz36 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К4']}балла.") 
abz37 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К4']}балла.")

table4 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table4.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table4.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table4.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 5. Удовлетворенность условиями оказания услуг")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz40 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz41 = otchet.add_paragraph("Показатель 5.1. Доля получателей услуг, которые готовы рекомендовать учреждение образования родственникам и знакомым")
abz42 = otchet.add_paragraph("Показатель 5.2. Доля получателей услуг, удовлетворенных организационными условиями предоставления услуг (графиком и режимом работы образовательного учреждения) (в % от общего числа опрошенных получателей услуг)")
abz43 = otchet.add_paragraph("Показатель 5.3. Доля получателей услуг, удовлетворенных в целом условиями оказания услуг в учреждении образования (в % от общего числа опрошенных получателей услуг).")
abz44 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table14 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Преком', 'Порг.услуд', 'Пуд', 'К5']]
min_value = table14['К5'].min()
max_value = table14['К5'].max()
mean_value = table14['К5'].mean()
sorted_table = table14.sort_values(by='К5', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz45 = otchet.add_paragraph(f"Итоговые баллы по критерию «Удовлетворенность условиями оказания услуг» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz46 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К5']}балла.") 
abz47 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К5']}балла.")

table5 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table5.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table5.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table5.cell(i+1, j).text = str(sorted_table.values[i, j])

under_zag = otchet.add_paragraph()
run = under_zag.add_run("Итоговая оценка качества условий оказания услуг образовательными учреждениями. Рейтинг учреждений")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Общий рейтинг образовательных учреждений.")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

table15 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'К5', 'К5', 'К5', 'К5', 'К5', 'Общий балл']]
min_value = table15['Общий балл'].min()
max_value = table15['Общий балл'].max()
mean_value = table15['Общий балл'].mean()
sorted_table = table15.sort_values(by='Общий балл', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(1)

sorted_table['рейтинг'] = range(1, len(sorted_table) + 1)

table6 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table6.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table6.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table6.cell(i+1, j).text = str(sorted_table.values[i, j])

abz45 = otchet.add_paragraph(f"Итоговый анализ и оценка качества работы образовательных учреждений позволяет определить лучшие учреждения по результатам мониторинга.  Общий балл организаций варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по сумме критериев {mean_value}.")
abz46 = otchet.add_paragraph(f"Среди образовательных учреждений {plase} в первую тройку лидеров вошли следующие учреждения:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['Общий балл']}балла.") 
abz47 = otchet.add_paragraph("Последнюю строку рейтинга занимает")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['Общий балл']}балла.")

under_zag = otchet.add_paragraph()
run = under_zag.add_run(f"Основные выводы и рекомендации по результатам независимой оценки качества условий оказания услуг образовательных учреждений {plase}.")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
under_zag = otchet.add_paragraph()
run = under_zag.add_run(f"Основные выводы по результатам независимой оценки.")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

list_krit = {'Наименование критерия':['Критерий 1', 'Критерий 2', 'Критерий 3', 'Критерий 4', 'Критерий 5'],
             'Значение':[table['К1'].mean(), table11['К2'].mean(), table12['К3'].mean(), table13['К4'].mean(), table14['К5'].mean()]}
list_krit = pd.DataFrame(list_krit)
sorted_list_krit = list_krit.sort_values(by='Значение')
sorted_list_krit = sorted_list_krit.reset_index(drop=True)

# Создаем словарь для массовой замены значений
list_krit_dict = {'Критерий 1': '1. Открытость и доступность информации', 
                  'Критерий 2': '2. Комфортность условий предоставления услуг', 
                  'Критерий 3': '3. Доступность услуг для инвалидов', 
                  'Критерий 4': '4. Доброжелательность, вежливость работников организации', 
                  'Критерий 5': '5. Удовлетворенность условиями оказания услуг'
                 }
# Массово заменяем значения в столбце 'Имя'
sorted_list_krit['Наименование критерия'] = sorted_list_krit['Наименование критерия'].replace(list_krit_dict)

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz50 = otchet.add_paragraph(f"Согласно результатам проведённого исследования, основным недостатком у данных учреждений является {sorted_list_krit.at[0, 'Наименование критерия']}. ")
abz51 = otchet.add_paragraph(f"Также есть проблемы с {sorted_list_krit.at[1, 'Наименование критерия']} и {sorted_list_krit.at[2, 'Наименование критерия']}. ")# Вставка графика в документ Word

# Создаем таблицу с нужным количеством строк и столбцов
table20 = otchet.add_table(rows=1, cols=4)
table20.style = 'Table Grid'  # Применяем стиль таблицы

# Заголовки столбцов
hdr_cells = table20.rows[0].cells
hdr_cells[0].text = 'Name_org'
hdr_cells[1].text = 'Балл'
hdr_cells[2].text = 'Рейтинг'
hdr_cells[3].text = 'Недостатки'

# Проходимся по каждой строке и добавляем данные в таблицу
for index, row in output_df.iterrows():
    row_cells = table20.add_row().cells
    row_cells[0].text = str(row['Name_org'])
    row_cells[1].text = str(Raschet_ballov.loc[Raschet_ballov['Наименование организации или П/Н по списку'] == row['Name_org'], 'Общий балл'].values[0])
    row_cells[2].text = str(sorted_table.loc[sorted_table['Наименование организации или П/Н по списку'] == row['Name_org'], 'рейтинг'].values[0])
    row_cells[3].text = f"Недостатки на стенде: {str(row['bad_stend'])}\n"\
                        f"Недостатки на сайте: {str(row['bad_sait'])}\n"\
                        f"Недостатки функционирование дистанционных способов связи: {str(row['bad_dist'])}\n"\
                        f"Недостатки комфортности условий предоставления услуг: {str(row['bad_komf'])}\n"\
                        f"Недостатки в разрезе оборудования для инвалидов: {str(row['obor_inv'])}\n"\
                        f"Недостатки доступности среды для инвалидов: {str(row['sreda_inv'])}"





button = st.button("получить готовый файл расчет баллов")

if button:
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    Raschet_ballov.to_excel(writer, index=False, sheet_name='Sheet1')
    writer.close()
    output.seek(0)
    st.download_button(
        label="Загрузить",
        data=output,
        file_name='результаты.xlsx',
        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


button = st.button("Получить готовый отчет")
if button:
    bio = io.BytesIO()
    otchet.save(bio)
    st.download_button(
        label="Скачать",
        data=bio.getvalue(),
        file_name="Отчет.docx",
        mime="docx"
    )
