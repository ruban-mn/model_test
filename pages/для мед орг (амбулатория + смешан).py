import pandas as pd
import numpy as np
import openpyxl
import docx
import streamlit as st
import io
import time as tm
import re
from docx import Document
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import math



st.subheader('Независимая оценка качества услуг медицинских организаций')

uploaded_file = st.file_uploader("**Загрузите файл сводную по чек-листам амбулатория**", type=["xls", "xlsx"])
chek_list = pd.read_excel(uploaded_file)

uploaded_file2 = st.file_uploader("**Загрузите файл с массивом анкет амбулатория**", type=["xls", "xlsx"])
Answers_respond = pd.read_excel(uploaded_file2)

uploaded_file3 = st.file_uploader("**Загрузите файл с массивом анкет амбулатория+стационар**", type=["xls", "xlsx"])
Answers_respond1 = pd.read_excel(uploaded_file3)

#uploaded_file4 = st.file_uploader("**Загрузите файл с массивом анкет стационар**", type=["xls", "xlsx"])
#Answers_respond2 = pd.read_excel(uploaded_file4)

plase = st.text_input('Введите название территории в родительном падеже', 'н-р Московской области')

## подгружаем сводную по чек листам
#chek_list=pd.read_excel(r"C:\отчеты 2024\2024 14 Ставрополь нок мед\НОКМ Ставрополь ЧекЛист 2024.xlsx")
##подгружаем массив с ответами респондентов
#Answers_respond=pd.read_excel(r"C:\отчеты 2024\2024 14 Ставрополь нок мед\Анкета АМБ Ставрополь 2024.xlsx")
#Answers_respond1=pd.read_excel(r"C:\отчеты 2024\2024 14 Ставрополь нок мед\Анкета СТАЦ АМБ Ставрополь 2024.xlsx")
tm.sleep(30)

Answers_respond_list = Answers_respond.columns.tolist() ##извлекаем наименования столбцов в список

New_col = []  # Создаем пустой список
for i in range(29):  # Цикл от 0 до 18
    sim = i   # присваиваем номер
    New_col.append('v' + str(sim))  # добавляем новый номер вопрса в список

dictionary = dict(zip(Answers_respond_list, New_col)) # создаем  словарь для переименования стобцов
Answers_respond = Answers_respond.rename(columns=dictionary) # переименовываем столбцы в начальном датафрейме
Answers_respond['v30'] = Answers_respond['v8'].map({'менее 1 часа': 1, '3 часа': 3, '24 часа и более': 25, '12 часов': 12})
Answers_respond['v31'] = Answers_respond['v22'].map({'менее 7 календарных дней': 6, '7 календарных дней': 7, '10 календарных дней': 10, '12 календарных дней': 12, '13 календарных дней': 13, '14 календарных дней и более': 15})

# Рассчитываем значение для нового столбца
def calculate_value(row):
    total_answers = sum([val for val in row['v30'] if not math.isnan(val)])

    result_sum = sum([(i+1)*val for i, val in enumerate(row['v30']) if not math.isnan(val)])
    result = round(result_sum / total_answers)
    
    if result == 24:
        return 10
    elif result == 23:
        return 20
    elif result == 22:
        return 40
    elif 21 <= result <= 13:
        return 60
    elif result > 24:
        return 1
    else:
        return 100


result_df = Answers_respond.groupby('v0').apply(calculate_value).reset_index()
result_df.columns = ['v0', 'v30']

# Рассчитываем значение для нового столбца
def calculate_value1(row):
    total_answers = sum([val for val in row['v31'] if not math.isnan(val)])

    result_sum = sum([(i+1)*val for i, val in enumerate(row['v31']) if not math.isnan(val)])
    result = round(result_sum / total_answers)
    
    if result == 14:
        return 10
    elif result == 13:
        return 20
    elif result == 12:
        return 40
    elif 11 <= result <= 8:
        return 60
    elif result > 14:
        return 1
    else:
        return 100


result_df1 = Answers_respond.groupby('v0').apply(calculate_value1).reset_index()
result_df1.columns = ['v0', 'v31']

# Создание нового DataFrame для хранения результатов подсчета, считам количество ответов да на вопросы анкеты
ans_res = pd.DataFrame({'v0': Answers_respond['v0'].unique()})



selected_columns = ['v1', 'v2', 'v3', 'v4', 'v5', 'v11', 'v13', 'v14', 'v15', 'v19', 'v23', 'v24', 'v25', 'v26', 'v27', 'v28', 'v30', 'v31']
# Используем цикл для подсчета значений и создания новых столбцов
for col in selected_columns:
    value = 'да'  # Значение, которое мы считаем
    count_col_name = f'_{col}_'
    counts = Answers_respond[Answers_respond[col] == value].groupby('v0').size().reset_index(name=count_col_name)
    ans_res = ans_res.merge(counts, on='v0', how='left')

ans_res = ans_res.merge(result_df, on='v0', how='left')

ans_res = ans_res.merge(result_df1, on='v0', how='left')                                                                      

ans_res['_v40_'] = (ans_res['v30'] + ans_res['v31'])/2

ans_res['_v41_'] = (ans_res['_v13_'] + ans_res['_v23_'])/2

ans_res.rename(columns={'_v5_': '_v42_', '_v15_': '_v43_', '_v19_': '_v44_', '_v11_': '_v45_', '_v14_': '_v46_', '_v24_': '_v47_', '_v25_': '_v48_', '_v26_': '_v49_', '_v27_': '_v50_', '_v28_': '_v51_'}, inplace=True)
    
ans_res = ans_res.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
ans_res = ans_res.reset_index(drop=True)

answer_amb = ans_res.filter(items=['v0','_v1_', '_v2_', '_v3_', '_v4_', '_v40_', '_v41_', '_v42_', '_v43_', '_v44_', '_v45_', '_v46_', '_v47_', '_v48_', '_v49_', '_v50_', '_v51_'])

count_series = Answers_respond.groupby('v0')['v13'].count().reset_index()
count_series.columns = ['v0', '_v1_']
count_series1 = Answers_respond.groupby('v0')['v23'].count().reset_index()
count_series1.columns = ['v0', '_v2_']
answer_amb['_v52_'] = (count_series['_v1_'] + count_series1['_v2_'])/2
count_series2 = Answers_respond.groupby('v0')['v5'].count().reset_index()
count_series2.columns = ['v0', '_v3_']
answer_amb['_v53_'] = count_series2['_v3_']
count_series3 = Answers_respond.groupby('v0')['v11'].count().reset_index()
count_series3.columns = ['v0', '_v4_']
answer_amb['_v54_'] = count_series3['_v4_']
count_series4 = Answers_respond.groupby('v0')['v14'].count().reset_index()
count_series4.columns = ['v0', '_v5_']
answer_amb['_v55_'] = count_series4['_v5_']
answer_amb['_v56_'] = answer_amb['_v55_']
# то же самое для смешанной анкеты


Answers_respond_list1 = Answers_respond1.columns.tolist() ##извлекаем наименования столбцов в список

New_col1 = []  # Создаем пустой список
for i in range(36):  # Цикл от 0 до 18
    sim = i   # присваиваем номер
    New_col1.append('v' + str(sim))  # добавляем новый номер вопрса в список

dictionary1 = dict(zip(Answers_respond_list1, New_col1)) # создаем  словарь для переименования стобцов
Answers_respond1 = Answers_respond1.rename(columns=dictionary1) # переименовываем столбцы в начальном датафрейме
Answers_respond1['v36'] = Answers_respond1['v8'].map({'менее 15 календарных дней': 14, '15 календарных дней': 15})
Answers_respond1['v37'] = Answers_respond1['v15'].map({'менее 1 часа': 1, '3 часа': 3, '6 часов': 6})
Answers_respond1['v38'] = Answers_respond1['v29'].map({'менее 7 календарных дней': 6, '7 календарных дней': 7, '10 календарных дней': 10, '14 календарных дней': 14})

# Рассчитываем значение для нового столбца
def calculate_value(row):
    total_answers = sum([val for val in row['v36'] if not math.isnan(val)])

    result_sum = sum([(i+1)*val for i, val in enumerate(row['v36']) if not math.isnan(val)])
    result = round(result_sum / total_answers)
    
    if result == 14:
        return 10
    elif result == 13:
        return 20
    elif result == 12:
        return 40
    elif 11 <= result <= 8:
        return 60
    else:
        return 100


result_df = Answers_respond1.groupby('v0').apply(calculate_value).reset_index()
result_df.columns = ['v0', 'v36']

# Рассчитываем значение для нового столбца
def calculate_value1(row):
    total_answers = sum([val for val in row['v37'] if not math.isnan(val)])

    result_sum = sum([(i+1)*val for i, val in enumerate(row['v37']) if not math.isnan(val)])
    result = round(result_sum / total_answers)
    
    if result == 14:
        return 10
    elif result == 13:
        return 20
    elif result == 12:
        return 40
    elif 11 <= result <= 8:
        return 60
    else:
        return 100


result_df1 = Answers_respond1.groupby('v0').apply(calculate_value1).reset_index()
result_df1.columns = ['v0', 'v37']

# Рассчитываем значение для нового столбца
def calculate_value2(row):
    total_answers = sum([val for val in row['v38'] if not math.isnan(val)])

    result_sum = sum([(i+1)*val for i, val in enumerate(row['v38']) if not math.isnan(val)])
    result = round(result_sum / total_answers)
    
    if result == 14:
        return 10
    elif result == 13:
        return 20
    elif result == 12:
        return 40
    elif 11 <= result <= 8:
        return 60
    else:
        return 100


result_df2 = Answers_respond1.groupby('v0').apply(calculate_value2).reset_index()
result_df2.columns = ['v0', 'v38']

# Создание нового DataFrame для хранения результатов подсчета, считам количество ответов да на вопросы анкеты
ans_res1 = pd.DataFrame({'v0': Answers_respond1['v0'].unique()})

selected_columns = ['v1', 'v2', 'v3', 'v4', 'v5', 'v9', 'v11', 'v13', 'v14', 'v18', 'v20', 'v21', 'v22', 'v26', 'v30', 'v31', 'v32', 'v33', 'v34', 'v35', 'v36', 'v37', 'v38']
# Используем цикл для подсчета значений и создания новых столбцов
for col in selected_columns:
    value = 'да'  # Значение, которое мы считаем
    count_col_name = f'_{col}_'
    counts = Answers_respond1[Answers_respond1[col] == value].groupby('v0').size().reset_index(name=count_col_name)
    ans_res1 = ans_res1.merge(counts, on='v0', how='left')

ans_res1 = ans_res1.merge(result_df, on='v0', how='left')

ans_res1 = ans_res1.merge(result_df1, on='v0', how='left')

ans_res1 = ans_res1.merge(result_df2, on='v0', how='left')

ans_res1['_v40_'] = (ans_res1['v36'] + ans_res1['v37'] + ans_res1['v38'])/3

ans_res1['_v41_'] = round((ans_res1['_v9_'] + ans_res1['_v20_'] + ans_res1['_v30_'])/3, 2)

ans_res1['_v42_'] = (ans_res1['_v5_'] + ans_res1['_v11_'])/2

ans_res1.rename(columns={'_v22_': '_v43_', '_v26_': '_v44_'}, inplace=True)

ans_res1['_v45_'] = (ans_res1['_v13_'] + ans_res1['_v18_'])/2

ans_res1['_v46_'] = (ans_res1['_v14_'] + ans_res1['_v21_'])/2

ans_res1.rename(columns={'_v31_': '_v47_', '_v32_': '_v48_', '_v33_': '_v49_', '_v34_': '_v50_', '_v35_': '_v51_'}, inplace=True)

ans_res1 = ans_res1.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
ans_res1 = ans_res1.reset_index(drop=True)

answer_amb_stat = ans_res1.filter(items=['v0','_v1_', '_v2_', '_v3_', '_v4_', '_v40_', '_v41_', '_v42_', '_v43_', '_v44_', '_v45_', '_v46_', '_v47_', '_v48_', '_v49_', '_v50_', '_v51_'])

count_series5 = Answers_respond1.groupby('v0')['v9'].count().reset_index()
count_series5.columns = ['v0', '_v1_']
count_series6 = Answers_respond1.groupby('v0')['v20'].count().reset_index()
count_series6.columns = ['v0', '_v2_']
count_series7 = Answers_respond1.groupby('v0')['v30'].count().reset_index()
count_series7.columns = ['v0', '_v3_']
answer_amb_stat['_v52_'] = round((count_series5['_v1_'] + count_series6['_v2_'] + + count_series7['_v3_'])/3, 2)
count_series8 = Answers_respond1.groupby('v0')['v5'].count().reset_index()
count_series8.columns = ['v0', '_v4_']
count_series9 = Answers_respond1.groupby('v0')['v11'].count().reset_index()
count_series9.columns = ['v0', '_v5_']
answer_amb_stat['_v53_'] = (count_series8['_v4_'] + count_series9['_v5_'])/2
count_series10 = Answers_respond1.groupby('v0')['v13'].count().reset_index()
count_series10.columns = ['v0', '_v6_']
count_series11 = Answers_respond1.groupby('v0')['v18'].count().reset_index()
count_series11.columns = ['v0', '_v7_']
answer_amb_stat['_v54_'] = (count_series10['_v6_'] + count_series11['_v7_'])/2
count_series12 = Answers_respond1.groupby('v0')['v14'].count().reset_index()
count_series12.columns = ['v0', '_v8_']
count_series13 = Answers_respond1.groupby('v0')['v21'].count().reset_index()
count_series13.columns = ['v0', '_v9_']
answer_amb_stat['_v55_'] = (count_series12['_v8_'] + count_series13['_v9_'])/2
answer_amb_stat['_v56_'] = count_series8['_v4_']


ans_all = pd.concat([answer_amb, answer_amb_stat], ignore_index=True)

ans_all = ans_all.sort_values(by='v0') # сортируем таблицу по возрастанию по столбцу наименования
ans_all = ans_all.reset_index(drop=True)

name_org = chek_list.filter(like='Наименование организации').copy()

Raschet_ballov = name_org
Raschet_ballov['Истенд'] = chek_list.filter(like='На СТЕНДЕ').sum(axis=1)
Raschet_ballov['Исайт'] = chek_list.filter(like='На САЙТЕ').sum(axis=1)
Raschet_ballov['Инорм-стенд'] = chek_list.filter(like='На СТЕНДЕ').count(axis=1)
Raschet_ballov['Инорм-сайт'] = chek_list.filter(like='На САЙТЕ').count(axis=1)
Raschet_ballov['Пинф'] = round(0.5*((Raschet_ballov['Истенд']/Raschet_ballov['Инорм-стенд'])+(Raschet_ballov['Исайт']/Raschet_ballov['Инорм-сайт']))*100, 2)
Raschet_ballov['Тдист'] = 30
Raschet_ballov['Сдист'] = chek_list.filter(like='Наличие и функционирование на официальном сайте').sum(axis=1)
Raschet_ballov['Пдист'] = Raschet_ballov['Тдист']*Raschet_ballov['Сдист']
Raschet_ballov['Пдист'].where(Raschet_ballov['Пдист'] <= 100, 100, inplace=True)
Raschet_ballov['Устенд'] = ans_all['_v2_']
Raschet_ballov['Усайт'] = ans_all['_v4_']
Raschet_ballov['Уобщ-стенд'] = ans_all['_v1_']
Raschet_ballov['Уобщ-сайт'] = ans_all['_v3_']
Raschet_ballov['Поткруд'] = round(0.5*((Raschet_ballov['Устенд']/Raschet_ballov['Уобщ-стенд'])+(Raschet_ballov['Усайт']/Raschet_ballov['Уобщ-сайт']))*100, 2)
Raschet_ballov['К1'] = round(0.3*Raschet_ballov['Пинф'] + 0.3*Raschet_ballov['Пдист'] + 0.4*Raschet_ballov['Поткруд'], 2)
Raschet_ballov['Ткомф'] = chek_list.filter(like='Обеспечение в организации комфортных условий').sum(axis=1)
Raschet_ballov['Скомф'] = 20
Raschet_ballov['Пкомф.усл'] = Raschet_ballov['Ткомф']*Raschet_ballov['Скомф']
Raschet_ballov['Пкомф.усл'].where(Raschet_ballov['Пкомф.усл'] <= 100, 100, inplace=True)
Raschet_ballov['ожид'] = ans_all['_v40_']
Raschet_ballov['Усвоевр'] = ans_all['_v41_']
Raschet_ballov['Чобщ'] = ans_all['_v52_']
Raschet_ballov['Пожид'] = (round(Raschet_ballov['Усвоевр']/Raschet_ballov['Чобщ']*100, 2) + Raschet_ballov['ожид'])/2
Raschet_ballov['Укомф'] = ans_all['_v42_']
Raschet_ballov['Чобщ0'] = ans_all['_v53_']
Raschet_ballov['Пкомфуд'] = round(Raschet_ballov['Укомф']/Raschet_ballov['Чобщ0']*100, 2)
Raschet_ballov['К2'] = round(0.3*Raschet_ballov['Пкомф.усл'] + 0.4*Raschet_ballov['Пожид'] + 0.3*Raschet_ballov['Пкомфуд'], 2)
Raschet_ballov['Торгдост'] = chek_list.filter(like='Оборудование территории').sum(axis=1)
Raschet_ballov['Соргдост'] = 20
Raschet_ballov['Поргдост'] = Raschet_ballov['Торгдост']*Raschet_ballov['Соргдост']
Raschet_ballov['Поргдост'].where(Raschet_ballov['Поргдост'] <= 100, 100, inplace=True)
Raschet_ballov['Туслугдост'] = chek_list.filter(like='Обеспечение в организации условий доступности').sum(axis=1)
Raschet_ballov['Суслугдост'] = 20
Raschet_ballov['Пуслугдост'] = Raschet_ballov['Туслугдост']*Raschet_ballov['Суслугдост']
Raschet_ballov['Пуслугдост'].where(Raschet_ballov['Пуслугдост'] <= 100, 100, inplace=True)
Raschet_ballov['Удост'] = ans_all['_v44_']
Raschet_ballov['Чинв'] = ans_all['_v43_']
Raschet_ballov['Пдостуд'] = round(Raschet_ballov['Удост']/Raschet_ballov['Чинв']*100, 2)
Raschet_ballov['К3'] = round(0.3*Raschet_ballov['Поргдост'] + 0.4*Raschet_ballov['Пуслугдост'] + 0.3*Raschet_ballov['Пдостуд'], 2)
Raschet_ballov['Уперв.конт'] = ans_all['_v45_']
Raschet_ballov['Чобщ1'] = ans_all['_v54_']
Raschet_ballov['Пперв.контуд'] = round(Raschet_ballov['Уперв.конт']/Raschet_ballov['Чобщ1']*100, 2)
Raschet_ballov['Уоказ.услуг'] = ans_all['_v46_']
Raschet_ballov['Чобщ2'] = ans_all['_v55_']
Raschet_ballov['Показ.услугуд'] = round(Raschet_ballov['Уоказ.услуг']/Raschet_ballov['Чобщ2']*100, 2)
Raschet_ballov['Увежл.дист'] = ans_all['_v48_']
Raschet_ballov['Чобщ_ус'] = ans_all['_v47_']
Raschet_ballov['Пвежл.дистуд'] = round(Raschet_ballov['Увежл.дист']/Raschet_ballov['Чобщ_ус']*100, 2)
Raschet_ballov['К4'] = round(0.4*Raschet_ballov['Пперв.контуд'] + 0.4*Raschet_ballov['Показ.услугуд'] + 0.2*Raschet_ballov['Пвежл.дистуд'], 2)
Raschet_ballov['Уреком'] = ans_all['_v49_']
Raschet_ballov['Чобщ3'] = ans_all['_v56_']
Raschet_ballov['Преком'] = round(Raschet_ballov['Уреком']/Raschet_ballov['Чобщ3']*100, 2)
Raschet_ballov['Уорг.усл'] = ans_all['_v50_']
Raschet_ballov['Чобщ4'] = ans_all['_v56_']
Raschet_ballov['Порг.услуд'] = round(Raschet_ballov['Уорг.усл']/Raschet_ballov['Чобщ4']*100, 2)
Raschet_ballov['Ууд'] = ans_all['_v51_']
Raschet_ballov['Чобщ5'] = ans_all['_v56_']
Raschet_ballov['Пуд'] = round(Raschet_ballov['Ууд']/Raschet_ballov['Чобщ5']*100, 2)
Raschet_ballov['К5'] = round(0.3*Raschet_ballov['Преком'] + 0.2*Raschet_ballov['Порг.услуд'] + 0.5*Raschet_ballov['Пуд'], 2)
Raschet_ballov['Общий балл'] = round((Raschet_ballov['К1']+Raschet_ballov['К2']+Raschet_ballov['К3']+Raschet_ballov['К4']+Raschet_ballov['К5'])/5, 2)



row_chek_list = chek_list.columns.tolist()

New_col_for_chek_list = []  # Создаем пустой список
for i in range(chek_list.shape[1]+1):  # Цикл от 0 до 18
    sim = i   # присваиваем номер
    New_col_for_chek_list.append('us' + str(sim))  # добавляем новый номер вопрса в список

dict_chek = dict(zip(row_chek_list, New_col_for_chek_list))
chek_list = chek_list.rename(columns=dict_chek) # переименовываем столбцы в начальном датафрейме


name_org1 = pd.DataFrame({'us0': chek_list['us0']}) 
chek_list_stend = chek_list.iloc[:, 1:60]  # Датафрейм с 1-5 столбцами
chek_list_sait = chek_list.iloc[:, 61:120]  # Датафрейм с 6-10 столбцами = df.iloc[:, 0:5]  # Датафрейм с 1-5 столбцами
chek_list_dist = chek_list.iloc[:, 121:124]  # Датафрейм с 6-10 столбцами
chek_list_komf = chek_list.iloc[:, 125:133]
chek_list_obor_inv = chek_list.iloc[:, 134:138]
chek_list_sreda_inv = chek_list.iloc[:, 139:144]

chek_list_stend = pd.concat([name_org1, chek_list_stend], axis=1)
chek_list_sait = pd.concat([name_org1, chek_list_sait], axis=1)
chek_list_dist = pd.concat([name_org1, chek_list_dist], axis=1)
chek_list_komf = pd.concat([name_org1, chek_list_komf], axis=1)
chek_list_obor_inv = pd.concat([name_org1, chek_list_obor_inv], axis=1)
chek_list_sreda_inv = pd.concat([name_org1, chek_list_sreda_inv], axis=1)

def process_row(row):
    first_column_value = row.iloc[0]
    row_subset = row.iloc[1:]

    zero_values = row_subset[row_subset == 0].index.tolist()  # Получаем список индексов столбцов с нулевыми значениями

    return pd.Series([first_column_value, zero_values], index=['First_Column_Value', 'Zero_Values'])
# Применяем функцию и конвертируем результаты в DataFrame
nedostatki_stend = chek_list_stend.apply(process_row, axis=1)
nedostatki_stend = nedostatki_stend[nedostatki_stend.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_stend = nedostatki_stend.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))

nedostatki = nedostatki_stend
nedostatki = nedostatki.rename(columns={'First_Column_Value':'Name_org', 'Zero_Values':'bad_stend'})

nedostatki_sait = chek_list_sait.apply(process_row, axis=1)
nedostatki_sait = nedostatki_sait[nedostatki_sait.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_sait = nedostatki_sait.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_sait'] = nedostatki_sait['Zero_Values']

nedostatki_dist = chek_list_dist.apply(process_row, axis=1)
nedostatki_dist = nedostatki_dist[nedostatki_dist.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_dist = nedostatki_dist.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_dist'] = nedostatki_dist['Zero_Values']

nedostatki_komf = chek_list_komf.apply(process_row, axis=1)
nedostatki_komf = nedostatki_komf[nedostatki_dist.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_komf = nedostatki_komf.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['bad_komf'] = nedostatki_komf['Zero_Values']

nedostatki_obor_inv = chek_list_obor_inv.apply(process_row, axis=1)
nedostatki_obor_inv = nedostatki_obor_inv[nedostatki_obor_inv.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_obor_inv = nedostatki_obor_inv.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['obor_inv'] = nedostatki_obor_inv['Zero_Values']

nedostatki_sreda_inv = chek_list_sreda_inv.apply(process_row, axis=1)
nedostatki_sreda_inv = nedostatki_sreda_inv[nedostatki_sreda_inv.apply(lambda x: len(x) > 0, axis=1)]
nedostatki_sreda_inv = nedostatki_sreda_inv.apply(lambda x: x.apply(lambda y: 'нет недостатков' if len(y) == 0 else y))
nedostatki['sreda_inv'] = nedostatki_sreda_inv['Zero_Values']

row_chek_list1 = row_chek_list.copy
new_list = [re.search(r'\[([^\]]+)\]', item).group(1) for item in row_chek_list[1:] if re.search(r'\[([^\]]+)\]', item)]
name_list = [row_chek_list[0]]
row_chek_list = name_list + new_list
dict_chek1 = dict(zip(New_col_for_chek_list, row_chek_list))

output_data = []

for index, row in nedostatki.iterrows():
    output_row = [row['Name_org']]
    for col in nedostatki.columns[1:]:
        if row[col] == "нет недостатков":
            output_row.append(row[col])
        else:
            values_to_find = row[col] 
            result_keys = str([dict_chek1[value] for value in values_to_find ])
            output_row.append(result_keys)
    output_data.append(output_row)

output_df = pd.DataFrame(output_data, columns=['Name_org', 'bad_stend', 'bad_sait', 'bad_dist', 'bad_komf', 'obor_inv', 'sreda_inv'])

otchet = Document()
section = otchet.sections[0]
section.left_margin = Cm(3.0)  # Левое поле 2 дюйма
section.right_margin = Cm(1.5)  # Правое поле 2 дюйма
section.top_margin = Cm(2.0)  # Верхнее поле 2 дюйма
section.bottom_margin = Cm(2.0)  # Нижнее поле 2 дюйма
# Устанавливаем шрифт по умолчанию для дальнейшего добавляемого текста
default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

# Добавляем параграф с измененными параметрами шрифта
zag = otchet.add_paragraph()
run = zag.add_run("Основные результаты исследования")
font = run.font
run.bold = True
font.size = Pt(18) 
zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Добавляем параграф с измененными параметрами шрифта
under_zag = otchet.add_paragraph()
run = under_zag.add_run("Результаты независимой оценки качества условий оказания услуг медицинскими учреждениями")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 1. Открытость и доступность информации об медицинском учреждении")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

# Добавляем текст с установленным шрифтом
abz =otchet.add_paragraph("Критерий представлен тремя показателями:")
abz1 = otchet.add_paragraph("Показатель 1.1.	 Соответствие информации о деятельности медицинского учреждения, размещенной на общедоступных информационных ресурсах, ее содержанию и порядку (форме), установленным нормативными правовыми актами (на информационных стендах в помещении медицинского учреждения; на официальном сайте медицинского учреждения в сети «Интернет»).")
abz2 = otchet.add_paragraph("Показатель 1.2. 	Наличие на официальном сайте медицинского учреждения информации о дистанционных способах обратной связи и взаимодействия с получателями услуг и их функционирование (абонентского номера телефона; адреса электронной почты; электронных сервисов (для подачи электронного обращения (жалобы, предложения), получения консультации по оказываемым услугам и иных.); раздела официального сайта «Часто задаваемые вопросы»; технической возможности выражения получателем услуг мнения о качестве условий оказания услуг медицинским учреждением (наличие анкеты для опроса граждан или гиперссылки на нее)).")
abz3 = otchet.add_paragraph("Показатель 1.3.	 Доля получателей услуг, удовлетворенных открытостью, полнотой и доступностью информации о деятельности медицинского учреждения, размещенной на информационных стендах в помещении медицинского учреждения, на официальном сайте медицинского учреждения в сети «Интернет» (в % от общего числа опрошенных получателей услуг).")
abz4 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz5 = otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пинф', 'Пдист', 'Поткруд', 'К1']]
min_value = table['К1'].min()
max_value = table['К1'].max()
mean_value = table['К1'].mean()
sorted_table = table.sort_values(by='К1', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz6 = otchet.add_paragraph(f"Итоговые баллы по критерию «Открытость и доступность информации о медицинском учреждении» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz7 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К1']}балла.")    
abz8 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
   otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К1']}балла.")

# Добавляем таблицу в документ
table1 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table1.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table1.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table1.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 2. Комфортность условий предоставления услуг")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz10 = otchet.add_paragraph("Критерий представлен двумя показателями:")
abz11 = otchet.add_paragraph("Показатель 2.1. Обеспечение в медицинском учреждении комфортных условий пребывания (транспортная/ пешая доступность медицинского учреждения, санитарное состояние помещений и территории учреждения, наличие и доступность питьевой воды, санитарно-гигиенических помещений, достаточность гардеробов)")
abz12 = otchet.add_paragraph("Показатель 2.3. Доля получателей услуг, удовлетворенных комфортностью предоставления услуг медицинским учреждением (в % от общего числа опрошенных получателей услуг).")
abz13 = otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table11 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пкомф.усл', 'Пожид', 'Пкомфуд', 'К2']]
min_value = table11['К2'].min()
max_value = table11['К2'].max()
mean_value = table11['К2'].mean()
sorted_table = table11.sort_values(by='К2', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz13_1 = otchet.add_paragraph(f"Итоговые баллы по критерию «Комфортность условий предоставления услуг» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz13_2 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К2']}балла.")    
abz13_3 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
   otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К2']}балла.")
    
# Добавляем таблицу в документ
table2 =otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table2.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table2.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table2.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 3. Доступность услуг для инвалидов")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz20 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz21 = otchet.add_paragraph("Показатель 3.1. Оборудование помещений медицинского учреждения и прилегающей к ней территории с учетом доступности для инвалидов (наличие оборудованных входных групп пандусами (подъемными платформами); наличие выделенных стоянок для автотранспортных средств инвалидов; наличие адаптированных лифтов, поручней, расширенных дверных проемов; наличие сменных кресел-колясок; наличие специально оборудованных санитарно-гигиенических помещений в учреждения социальной сферы).")
abz22 = otchet.add_paragraph("Показатель 3.2. Обеспечение в медицинском учреждении условий доступности, позволяющих инвалидам получать услуги наравне с другими (дублирование для инвалидов по слуху и зрению звуковой и зрительной информации; дублирование надписей, знаков и иной текстовой и графической информации знаками, выполненными рельефно-точечным шрифтом Брайля; возможность предоставления инвалидам по слуху (слуху и зрению) услуг сурдопереводчика (тифлосурдопереводчика); наличие альтернативной версии официального сайта учреждения социальной сферы в сети «Интернет» для инвалидов по зрению; помощь, оказываемая работниками учреждения социальной сферы, прошедшими необходимое обучение (инструктирование) по сопровождению инвалидов в помещениях учреждения социальной сферы и на прилегающей территории; наличие возможности предоставления услуги в дистанционном режиме или на дому). ")
abz23 = otchet.add_paragraph("Показатель 3.3. Доля получателей услуг, удовлетворенных доступностью услуг для инвалидов (в % от общего числа опрошенных получателей услуг – инвалидов)")
abz24 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table12 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Поргдост', 'Пуслугдост', 'Пдостуд', 'К3']]
min_value = table12['К3'].min()
max_value = table12['К3'].max()
mean_value = table12['К3'].mean()
sorted_table = table12.sort_values(by='К3', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz25 = otchet.add_paragraph(f"Итоговые баллы по критерию «Доступность услуг для инвалидов» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz26 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К3']}балла.") 
abz27 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К3']}балла.")

table3 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table3.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table3.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table3.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 4. Доброжелательность, вежливость работников учреждения")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz30 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz31 = otchet.add_paragraph("Показатель 4.1. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников медицинского учреждения, обеспечивающих первичный контакт и информирование получателя услуги при непосредственном обращении в организацию образования (в % от общего числа опрошенных получателей услуг)")
abz32 = otchet.add_paragraph("Показатель 4.2. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников медицинского учреждения, обеспечивающих непосредственное оказание услуги при обращении в организацию образования (в % от общего числа опрошенных получателей услуг)")
abz33 = otchet.add_paragraph("Показатель 4.3. Доля получателей услуг, удовлетворенных доброжелательностью, вежливостью работников медицинского учреждения при использовании дистанционных форм взаимодействия (в % от общего числа опрошенных получателей услуг).")
abz34 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table13 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Пперв.контуд', 'Показ.услугуд', 'Пвежл.дистуд', 'К4']]
min_value = table13['К4'].min()
max_value = table13['К4'].max()
mean_value = table13['К4'].mean()
sorted_table = table13.sort_values(by='К4', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz35 = otchet.add_paragraph(f"Итоговые баллы по критерию «Доброжелательность, вежливость работников учреждения» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz36 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К4']}балла.") 
abz37 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К4']}балла.")

table4 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table4.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table4.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table4.cell(i+1, j).text = str(sorted_table.values[i, j])

name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Критерий 5. Удовлетворенность условиями оказания услуг")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz40 = otchet.add_paragraph("Критерий представлен тремя показателями:")
abz41 = otchet.add_paragraph("Показатель 5.1. Доля получателей услуг, которые готовы рекомендовать социальное учреждение родственникам и знакомым")
abz42 = otchet.add_paragraph("Показатель 5.2. Доля получателей услуг, удовлетворенных организационными условиями предоставления услуг (графиком и режимом работы социального учреждения) (в % от общего числа опрошенных получателей услуг)")
abz43 = otchet.add_paragraph("Показатель 5.3. Доля получателей услуг, удовлетворенных в целом условиями оказания услуг в медицинском учреждении (в % от общего числа опрошенных получателей услуг).")
abz44 =otchet.add_paragraph("Максимальное количество баллов по данному критерию – 100,00.")

table14 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'Преком', 'Порг.услуд', 'Пуд', 'К5']]
min_value = table14['К5'].min()
max_value = table14['К5'].max()
mean_value = table14['К5'].mean()
sorted_table = table14.sort_values(by='К5', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(3)

abz45 = otchet.add_paragraph(f"Итоговые баллы по критерию «Удовлетворенность условиями оказания услуг» варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по критерию {mean_value}.")
abz46 = otchet.add_paragraph("Первые три лучших результата у организаций:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К5']}балла.") 
abz47 = otchet.add_paragraph("Три последних результата у организаций:")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['К5']}балла.")

table5 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table5.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table5.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table5.cell(i+1, j).text = str(sorted_table.values[i, j])

under_zag = otchet.add_paragraph()
run = under_zag.add_run("Итоговая оценка качества условий оказания услуг медицинскими учреждениями. Рейтинг учреждений")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
name_otchet = otchet.add_paragraph()
run = name_otchet.add_run("Общий рейтинг медицинских учреждений.")
run.bold = True
font = run.font
font.size = Pt(16) 
name_otchet.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

table15 = Raschet_ballov.loc[:, ['Наименование организации или П/Н по списку', 'К1', 'К2', 'К3', 'К4', 'К5', 'Общий балл']]
min_value = table15['Общий балл'].min()
max_value = table15['Общий балл'].max()
mean_value = table15['Общий балл'].mean()
sorted_table = table15.sort_values(by='Общий балл', ascending=False)
top_3_rows = sorted_table.head(3)
bad_3_rows = sorted_table.tail(1)

sorted_table['рейтинг'] = range(1, len(sorted_table) + 1)

table6 = otchet.add_table(sorted_table.shape[0]+1, sorted_table.shape[1])
table6.style = 'Table Grid'  # Применяем стиль таблицы
# Заголовки столбцов
for j in range(sorted_table.shape[-1]):
    table6.cell(0, j).text = sorted_table.columns[j]

# Данные из DataFrame
for i in range(sorted_table.shape[0]):
    for j in range(sorted_table.shape[-1]):
        table6.cell(i+1, j).text = str(sorted_table.values[i, j])

abz45 = otchet.add_paragraph(f"Итоговый анализ и оценка качества работы социальных учреждений позволяет определить лучшие учреждения по результатам мониторинга.  Общий балл организаций варьируются от {min_value} до {max_value} баллов. Средний итоговый балл по сумме критериев {mean_value}.")
abz46 = otchet.add_paragraph(f"Среди социальных учреждений {plase} в первую тройку лидеров вошли следующие учреждения:")
for index, row in top_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['Общий балл']}балла.") 
abz47 = otchet.add_paragraph("Последнюю строку рейтинга занимает")
for index, row in bad_3_rows.iterrows():
    otchet.add_paragraph(f"{row['Наименование организации или П/Н по списку']}, {row['Общий балл']}балла.")

under_zag = otchet.add_paragraph()
run = under_zag.add_run(f"Основные выводы и рекомендации по результатам независимой оценки качества условий оказания услуг медицинских учреждений {plase}.")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
under_zag = otchet.add_paragraph()
run = under_zag.add_run(f"Основные выводы по результатам независимой оценки.")
font = run.font
run.bold = True
font.size = Pt(16) 
under_zag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

list_krit = {'Наименование критерия':['Критерий 1', 'Критерий 2', 'Критерий 3', 'Критерий 4', 'Критерий 5'],
             'Значение':[table['К1'].mean(), table11['К2'].mean(), table12['К3'].mean(), table13['К4'].mean(), table14['К5'].mean()]}
list_krit = pd.DataFrame(list_krit)
sorted_list_krit = list_krit.sort_values(by='Значение')
sorted_list_krit = sorted_list_krit.reset_index(drop=True)

# Создаем словарь для массовой замены значений
list_krit_dict = {'Критерий 1': '1. Открытость и доступность информации', 
                  'Критерий 2': '2. Комфортность условий предоставления услуг', 
                  'Критерий 3': '3. Доступность услуг для инвалидов', 
                  'Критерий 4': '4. Доброжелательность, вежливость работников организации', 
                  'Критерий 5': '5. Удовлетворенность условиями оказания услуг'
                 }
# Массово заменяем значения в столбце 'Имя'
sorted_list_krit['Наименование критерия'] = sorted_list_krit['Наименование критерия'].replace(list_krit_dict)

# Добавляем таблицу в документ
table100 = otchet.add_table(rows=1, cols=2)  # Создаем таблицу с 1 строкой и 2 колонками

# Добавляем заголовки столбцов
hdr_cells = table100.rows[0].cells
hdr_cells[0].text = 'Наименование критерия'
hdr_cells[1].text = 'Среднее значение'

# Заполняем таблицу данными
for index, row in sorted_list_krit.iterrows():
    row_cells = table100.add_row().cells
    row_cells[0].text = row['Наименование критерия']
    row_cells[1].text = str(row['Значение'])

default_font = otchet.styles['Normal'].font
default_font.name = 'Times New Roman'
default_font.size = Pt(14)

abz50 = otchet.add_paragraph(f"Согласно результатам проведённого исследования, основным недостатком у данных учреждений является {sorted_list_krit.at[0, 'Наименование критерия']}. ")
abz51 = otchet.add_paragraph(f"Также есть проблемы с {sorted_list_krit.at[1, 'Наименование критерия']} и {sorted_list_krit.at[2, 'Наименование критерия']}. ")# Вставка графика в документ Word

if not Raschet_ballov.empty:
    print("DataFrame не является пустым")
else:
    print("DataFrame пустой")

if not sorted_table.empty:
    print("DataFrame не является пустым")
else:
    print("DataFrame пустой")

# Создаем таблицу с нужным количеством строк и столбцов
table20 = otchet.add_table(rows=1, cols=4)
table20.style = 'Table Grid'  # Применяем стиль таблицы

# Заголовки столбцов
hdr_cells = table20.rows[0].cells
hdr_cells[0].text = 'Наименование организации'
hdr_cells[1].text = 'Балл'
hdr_cells[2].text = 'Рейтинг'
hdr_cells[3].text = 'Недостатки'

# Проходимся по каждой строке и добавляем данные в таблицу
for index, row in output_df.iterrows():
    row_cells = table20.add_row().cells
    row_cells[0].text = str(row['Name_org'])
    row_cells[1].text = str(Raschet_ballov.loc[Raschet_ballov['Наименование организации или П/Н по списку'] == row['Name_org'], 'Общий балл'].values[0])
    row_cells[2].text = str(sorted_table.loc[sorted_table['Наименование организации или П/Н по списку'] == row['Name_org'], 'рейтинг'].values[0])

    row_cells[3].text = f"Недостатки на стенде: отсутствуют документы о {str(row['bad_stend'])}\n"\
                        f"Недостатки на сайте: отсутствуют документы о {str(row['bad_sait'])}\n"\
                        f"Недостатки функционирование дистанционных способов связи: {str(row['bad_dist'])}\n"\
                        f"Недостатки комфортности условий предоставления услуг: {str(row['bad_komf'])}\n"\
                        f"Недостатки в разрезе оборудования для инвалидов: {str(row['obor_inv'])}\n"\
                        f"Недостатки доступности среды для инвалидов: {str(row['sreda_inv'])}"


abz60 = otchet.add_paragraph(f"Рекомендации для организаций")
# Создаем таблицу с нужным количеством строк и столбцов
table30 = otchet.add_table(rows=1, cols=2)
table30.style = 'Table Grid'  # Применяем стиль таблицы

# Заголовки столбцов
hdr_cells = table30.rows[0].cells
hdr_cells[0].text = 'Наименование организации'
hdr_cells[1].text = 'Рекомендации'

# Проходимся по каждой строке и добавляем данные в таблицу
for index, row in output_df.iterrows():
    row_cells = table30.add_row().cells
    row_cells[0].text = str(row['Name_org'])
    
    K3_list = table15.loc[table15['Наименование организации или П/Н по списку'] == row['Name_org'], 'К3'].values[0]
        # Добавляем результат проверки условия в столбец "Недостатки"
    if K3_list < 90:
        result_k3 = ''.join(("Провести инструктаж и обучение сотрудников по вопросам обеспечения доступности организации для инвалидов, в т.ч. ",
                             "по сопровождению маломобильных граждан в помещениях организации и на прилегающей территории с учетом имеющихся у ",
                             "них ограничений и расстройств, по оказанию содействия инвалиду в беспрепятственном получении услуги.",
                             "определить и закрепить в должностных инструкциях и распорядительных документах ответственность конкретных работников ",
                             "(категорий работников) организации по сопровождению маломобильных граждан и оказанию им содействия в получении услуги.",
                             "Провести оптимизацию расположения кабинетов организации с учетом обеспечения их доступности для маломобильных граждан, ",
                             "перенести наиболее востребованные кабинеты и услуги на первый этаж"))
    else:
        result_k3 = 'нет рекомендаций'
    # Добавляем результат проверки условия в столбец "Недостатки"

    K4_list = table15.loc[table15['Наименование организации или П/Н по списку'] == row['Name_org'], 'К4'].values[0]
    
    if K4_list < 90:
        result_k4 = ''.join(("Провести обучение персонала организации по вопросам этики и деонтологии",
                             "Провести обучение персонала организации по вопросам клиентоцентричности, реализации стандарта «Государство для людей» ",
                             "и принципов проактивного предоставления услуг. Ввести на регулярной основе (не реже 1 раза в квартал) рабочие совещания ",
                             "с коллективом организации по вопросам соблюдения норм профессиональной этики и правил служебного поведения",
                             "Разработка внутреннего стандарта и памятки по предоставлению информации по телефону для специалистов организации",
                             "Внести изменения в трудовой договор с работниками организации, включив результаты НОК (отдельные критерии/показатели) ",
                             "в состав показателей результативности каждого работника, осуществляющего взаимодействие с получателями услуг, ",
                             "для установления стимулирующих выплат и использования в рамках программ нематериальной мотивации"))
    else:
        result_k4 = 'нет рекомендаций'    
    # Добавляем результат проверки условия в столбец "Недостатки"
    K5_list = table15.loc[table15['Наименование организации или П/Н по списку'] == row['Name_org'], 'К5'].values[0]
    
    if K5_list < 90:
        result_k5 = ''.join(("Провести внутренний аудит системы менеджмента качества в структурных подразделениях организации ",
                             "с целью реализации превентивных мер, направленных на совершенствование условий оказания услуг по перечню недостатков, ",
                             "выявленных в других организациях или структурных подразделениях (управление по прецедентам)",
                             "Провести анкетирование среди получателей социальных услуг, направленное на выявление глубинных причин ",
                             "неудовлетворенности условиями оказания услуг. Разработать карты клиентского пути, направленные на уточнение ",
                             "и оптимизацию сценариев и условий оказания услуг в организации. Разработать и провести информационную кампанию ",
                             "по информированию граждан о формах и видах оказываемых услуг, существующих преимуществах получения их в данной организации, ",
                             "возможных льготах и пр. (изготовление памяток, буклетов, написание статей в СМИ и т.п.) ",
                             "Внедрить в практику работы организации проекты, направленные на повышение эффективности предоставления социальных услуг ",
                             "и работу с целевыми категориями получателей услуг (Например, созданию условий для обучения людей с ментальными нарушениями, ",
                             "Территория возможностей для граждан серебряного возраста по организации и поддержке творческого досуга граждан пожилого возраста, ",
                             "удовлетворению и развитию их культурных, духовных потребностей). На уровне уполномоченного органа власти (учредителя организации) ",
                             "внести изменения в трудовой договор с руководителями организаций социальной сферы, включив общие результаты НОК ",
                             "и результаты исполнения планов по устранению недостатков, выявленных в ходе НОК в показатели результативности для установления ",
                             "стимулирующих выплат и использования в рамках программ нематериальной мотивации"))
    else:
        result_k5 = 'нет рекомендаций'
        
    row_cells[1].text = f"Разместить на стенде организации следующие документы: {str(row['bad_stend'])}\n"\
                        f"Разместить на сайте организации следующие документы: {str(row['bad_sait'])}\n"\
                        f"Обеспечить наличие в организации следующих видов дистанционного взаимодействия: {str(row['bad_dist'])}\n"\
                        f"Обеспечить следующие условия комфортности предоставления услуг: {str(row['bad_komf'])}\n"\
                        f"Обеспечить следующие оборудование для инвалидов: {str(row['obor_inv'])}\n"\
                        f"Обеспечить доступность среды для инвалидов, а именно: {str(row['sreda_inv'])}\n"\
                        f"В разрезе удовлетворенности доступностью услуг для инвалидов рекомендовано: {result_k3}\n"\
                        f"В разрезе удовлетворенности вежливостью и доброжелательностью работников учреждения рекомендовано:: {result_k4}\n"\
                        f"В разрезе удовлетворенности предоставлением услуг в целом рекомендовано: {result_k5}"



button = st.button("получить готовый файл расчет баллов")

if button:
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    Raschet_ballov.to_excel(writer, index=False, sheet_name='Sheet1')
    writer.close()
    output.seek(0)
    st.download_button(
        label="Загрузить",
        data=output,
        file_name='результаты.xlsx',
        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )



button = st.button("Получить готовый отчет")
if button:
    bio = io.BytesIO()
    otchet.save(bio)
    st.download_button(
        label="Скачать",
        data=bio.getvalue(),
        file_name="Отчет.docx",
        mime="docx"
    )
